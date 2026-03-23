#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script sinh tài liệu hướng dẫn DOCX chuyên nghiệp
cho dự án HR Analytics AI Team (11/02 - 16/02/2026)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_heading_with_color(doc, text, level, color=(0, 102, 204)):
    """Thêm heading với màu tùy chỉnh"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
    return heading

def add_table_border(table):
    """Thêm border cho table"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def set_cell_background(cell, color_code):
    """Đặt màu nền cho cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_code)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_pccv_doc():
    """Tạo tài liệu DOCX chuyên nghiệp"""
    doc = Document()
    
    # ===== COVER PAGE =====
    title = doc.add_heading('HƯỚNG DẪN THỰC HIỆN NHIỆM VỤ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('HR Analytics AI — Text-to-SQL Agent')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()  # Space
    
    # Thông tin cơ bản
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ('Mã học phần', '252BIM500601'),
        ('Đề tài', 'Ứng dụng AI trong phân tích rủi ro nghỉ việc nhân sự và hỗ trợ ra quyết định'),
        ('Hướng dẫn viên', 'TS. Trịnh Quang Việt'),
        ('Thời gian thực hiện', '11/02/2026 - 16/02/2026'),
        ('Team size', '5 người (Khải, Hân, Ninh, Gia, Uyên)'),
        ('Ngày phát hành', datetime.datetime.now().strftime('%d/%m/%Y'))
    ]
    
    for idx, (label, value) in enumerate(info_data):
        cells = info_table.rows[idx].cells
        cells[0].text = label
        cells[1].text = value
        set_cell_background(cells[0], 'E7E6E6')
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ===== MỤC LỤC =====
    add_heading_with_color(doc, 'MỤC LỤC', level=1)
    
    toc_items = [
        '1. Tổng quan dự án và kiến trúc hệ thống',
        '2. Link demo public và cách triển khai',
        '3. Hướng dẫn cài đặt và chạy dự án',
        '4. Bảng PCCV (Project Completion Control Version)',
        '5. Phân công chi tiết theo từng thành viên',
        '   5.1. Khải (K234060700) - Infrastructure & Connectivity',
        '   5.2. Hân (K234060691) - Semantic Layer & Vector Store',
        '   5.3. Ninh (K234060716) - Agentic Layer & Knowledge',
        '   5.4. Gia (K234060689) - Business Insights Analytics',
        '   5.5. Uyên (K234060737) - MLOps & Architecture',
        '6. Through-back chung (Q&A Showcase)',
        '7. Hướng dẫn cấu hình 5 nghiệp vụ HR trên UI',
        '8. Hướng dẫn test demo performance',
        '9. Yêu cầu đầu ra (Deliverables)',
        '10. Timeline, Checklist & Submission'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== PHẦN 1: TỔNG QUAN =====
    add_heading_with_color(doc, '1. Tổng quan dự án và kiến trúc hệ thống', level=1)
    
    add_heading_with_color(doc, '1.1. Mục tiêu cốt lõi', level=2, color=(0, 102, 204))
    p = doc.add_paragraph()
    p.add_run('Chuyển đổi HR từ thụ động → chủ động:').bold = True
    doc.add_paragraph('Sử dụng AI để phân tích rủi ro nghỉ việc và hỗ trợ HR Director ra quyết định chiến lược')
    doc.add_paragraph('Dân chủ hóa dữ liệu: Mọi nhân viên HR có thể tự hỏi câu hỏi bằng tiếng Việt tự nhiên (Text-to-SQL)')
    doc.add_paragraph('Bảo mật: Chỉ gửi Metadata cho LLM (Gemini), không gửi Raw Data của nhân viên')
    
    add_heading_with_color(doc, '1.2. Kiến trúc 4 Layer', level=2, color=(0, 102, 204))
    
    arch_table = doc.add_table(rows=5, cols=2)
    arch_table.style = 'Light Grid Accent 1'
    add_table_border(arch_table)
    
    arch_data = [
        ('Layer', 'Mô tả'),
        ('Data Layer', 'SQL Server 2019 + Views: init-db.sql, create_actionable_views.sql'),
        ('Analytics Layer', 'ML Random Forest: SMOTE, OOF, Feature Importance (Notebook)'),
        ('Agentic Layer', 'Wren AI: 29 pipelines (text-to-SQL, correction, intent, chart)'),
        ('Representation Layer', 'Wren UI: Next.js Chat, Modeling, Knowledge Management')
    ]
    
    for idx, (layer, desc) in enumerate(arch_data):
        cells = arch_table.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            cells[0].text = layer
            cells[1].text = desc
            set_cell_background(cells[0], 'E7E6E6')
    
    add_heading_with_color(doc, '1.3. 6 Docker Containers', level=2, color=(0, 102, 204))
    
    container_table = doc.add_table(rows=7, cols=4)
    container_table.style = 'Light Grid Accent 1'
    add_table_border(container_table)
    
    container_data = [
        ('Container', 'Port', 'Công nghệ', 'Vai trò'),
        ('wren-ui', '3000', 'Next.js 14.2', 'Giao diện chính: Chat, Modeling, Knowledge'),
        ('wren-ai-service', '5555', 'FastAPI', '29 AI pipelines + Gemini API integration'),
        ('wren-engine', '8080', 'MDL Compiler', 'Chuyển đổi MDL → SQL tối ưu'),
        ('ibis-server', '8000', 'Ibis Framework', 'Dịch SQL → T-SQL (cho MSSQL)'),
        ('qdrant', '6333', 'Vector DB', '6 collections: schema (52 docs), descriptions (17)'),
        ('hr-sql-server', '1433', 'SQL Server 2019', 'HR Analytics database (35 tables)')
    ]
    
    for idx, row_data in enumerate(container_data):
        cells = container_table.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            for cell_idx, value in enumerate(row_data):
                cells[cell_idx].text = value
    
    doc.add_paragraph()
    
    add_heading_with_color(doc, '1.4. LLM & Embedder Configuration', level=2, color=(0, 102, 204))
    
    config_items = [
        ('LLM Model', 'gemini/gemini-2.5-flash (1M token context, temp=0)'),
        ('Embedder', 'gemini-embedding-001 (768 dimensions)'),
        ('LLM Provider', 'LiteLLM (abstraction layer)'),
        ('Indexing', '29 AI pipelines (19 generation, 6 indexing, 4 retrieval)'),
        ('API Key', 'Tu nhap GEMINI_API_KEY trong file .env local (khong commit)')
    ]
    
    for label, value in config_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{label}: ').bold = True
        p.add_run(value)
    
    doc.add_page_break()
    
    # ===== PHẦN 2: DEMO PUBLIC =====
    add_heading_with_color(doc, '2. Link demo public và cách triển khai', level=1)
    
    add_heading_with_color(doc, '2.1. Public Link (Cloudflare Tunnel)', level=2, color=(0, 102, 204))
    
    p = doc.add_paragraph()
    p.add_run('Link demo hiện tại (thay đổi khi restart):').bold = True
    p = doc.add_paragraph()
    for run in p.runs:
        run.font.size = Pt(12)
    run = p.add_run('https://certification-lows-spy-tension.trycloudflare.com')
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Link này kết nối trực tiếp tới backend đang chạy, bao gồm:')
    
    features = [
        'Wren AI UI (Chat, Modeling, Knowledge)',
        'Wren AI Service (29 AI pipelines + Gemini API)',
        'Wren Engine + Ibis Server (SQL compilation + execution)',
        'MSSQL Server (HR Analytics database)',
        'Qdrant (Vector DB: 52 schema docs + 17 table descriptions)'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_with_color(doc, '2.2. Tạo lại Public Link (khi restart máy)', level=2, color=(0, 102, 204))
    
    steps = [
        'Khởi động Docker containers',
        'Chờ tất cả containers healthy (~30 giây)',
        'Tạo public tunnel bằng cloudflared',
        'Copy link mới và chia sẻ cho team'
    ]
    
    for idx, step in enumerate(steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)
    
    doc.add_page_break()
    
    # ===== PHẦN 3: CÀI ĐẶT =====
    add_heading_with_color(doc, '3. Hướng dẫn cài đặt và chạy dự án', level=1)
    
    add_heading_with_color(doc, '3.1. Yêu cầu hệ thống', level=2, color=(0, 102, 204))
    
    requirements = [
        'Windows 10/11 với Docker Desktop',
        'Docker Desktop (bật WSL2 integration)',
        'SQL Server 2019+ hoặc dùng container MSSQL',
        'Python 3.10+ (cho notebook)',
        'Git để clone repo'
    ]
    
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    add_heading_with_color(doc, '3.2. Clone repo từ GitLab', level=2, color=(0, 102, 204))
    
    p = doc.add_paragraph()
    p.add_run('git clone https://gitlab.com/boygia757-netizen/hr-ai-project.git').font.name = 'Courier New'
    p = doc.add_paragraph()
    p.add_run('cd hr-ai-project').font.name = 'Courier New'
    p = doc.add_paragraph()
    p.add_run('git checkout hr_domain_research').font.name = 'Courier New'
    
    add_heading_with_color(doc, '3.3. Khởi động Docker & Verify', level=2, color=(0, 102, 204))
    
    verify_steps = [
        'Chạy: docker compose up -d',
        'Chạy: docker ps (kiểm tra 6 containers UP)',
        'Health check: Invoke-RestMethod http://localhost:5555/health',
        'Mở UI: http://localhost:3000'
    ]
    
    for step in verify_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== PHẦN 4: BẢNG PCCV =====
    add_heading_with_color(doc, '4. Bảng PCCV (Project Completion Control Version)', level=1)
    
    doc.add_paragraph('Bảng PCCV liệt kê chi tiết 6 công việc chính của 5 thành viên trong 6 ngày (11-16/02/2026):')
    doc.add_paragraph()
    
    pccv_table = doc.add_table(rows=7, cols=6)
    pccv_table.style = 'Light Grid Accent 1'
    add_table_border(pccv_table)
    
    pccv_header = ['STT', 'Công việc', 'Người phụ trách', 'Source code', 'Deep dive Q', 'Deliverable']
    for idx, header in enumerate(pccv_header):
        cell = pccv_table.rows[0].cells[idx]
        cell.text = header
        set_cell_background(cell, '003366')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    
    pccv_data = [
        ('1', 'Infrastructure & Connectivity', 'Khải\n(K234060700)', '8 files', 'Q1-Q8', 'Docx + Diagram + 5 Use Cases'),
        ('2', 'Semantic Layer & Vector Store', 'Hân\n(K234060691)', '8 files', 'Q1-Q9', 'Docx + Diagram + 5 Use Cases'),
        ('3', 'Agentic Layer & Knowledge', 'Ninh\n(K234060716)', '11 files', 'Q1-Q9', 'Docx + Diagram + 5 Use Cases'),
        ('4', 'Business Analytics & Insights', 'Gia\n(K234060689)', '4 files', 'Q1-Q9', 'Docx + Report + Slides'),
        ('5', 'MLOps & Project Architecture', 'Uyên\n(K234060737)', '8 files', 'Q1-Q12', 'Docx + 3 Diagrams + Slides'),
        ('6', 'Through-back chung (Q&A)', 'Khải, Hân, Ninh', '-', '8 Q chung', 'Docx hoàn chỉnh Q&A')
    ]
    
    for idx, row_data in enumerate(pccv_data, 1):
        cells = pccv_table.rows[idx].cells
        for cell_idx, value in enumerate(row_data):
            cells[cell_idx].text = value
            for paragraph in cells[cell_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== PHẦN 5: PHÂN CÔNG CHI TIẾT =====
    add_heading_with_color(doc, '5. Phân công chi tiết theo từng thành viên', level=1)
    
    # ===== KHẢI =====
    add_heading_with_color(doc, '5.1. Khải (K234060700) — Infrastructure & Connectivity Owner', level=2)
    
    add_heading_with_color(doc, 'I. Công việc chính', level=3, color=(204, 102, 0))
    khải_jobs = [
        'Chạy lại toàn bộ dự án với cấu hình hiện tại, đảm bảo 6 containers hoạt động ổn định.',
        'Giải thích được vai trò từng container trong docker-compose.yaml: bootstrap, wren-engine, ibis-server, wren-ai-service, qdrant, wren-ui.',
        'Định nghĩa 5 nghiệp vụ HR mới (Modeling + Relationship + SQL Pair + Instruction) và demo thành công trên Wren AI UI.',
        'Show log realtime của 6 container khi 1 câu hỏi được gửi đi để chứng minh luồng dữ liệu.'
    ]
    for job in khải_jobs:
        doc.add_paragraph(job, style='List Number')
    
    add_heading_with_color(doc, 'II. Deep dive questions (trả lời trong Docx)', level=3, color=(204, 102, 0))
    
    khải_questions = [
        'Giải thích chi tiết luồng dữ liệu đi qua 6 container từ lúc user nhập câu hỏi đến lúc nhận kết quả?',
        'Tại sao cần Ibis Server? Ibis đóng vai trò gì giữa MDL và Native SQL của MSSQL?',
        'File ibisAdaptor.ts thực hiện những method nào để giao tiếp với Ibis container (query, dryPlan, metadata)?',
        'File wren.py trong providers/engine thực hiện kết nối tới Wren Engine bằng cách nào (GraphQL mutation PreviewSql)?',
        'Cấu hình .env gồm những biến môi trường nào? GEMINI_API_KEY được truyền vào container nào?',
        'Bootstrap container làm gì? Tại sao nó chỉ chạy 1 lần rồi dừng?',
        'Network "wren" trong docker-compose hoạt động thế nào để các container giao tiếp với nhau?',
        'Khi ibis-server gặp lỗi kết nối đến MSSQL, log hiển thị ở đâu và cách debug?'
    ]
    
    for idx, q in enumerate(khải_questions, 1):
        doc.add_paragraph(f'Q{idx}: {q}', style='List Bullet')
    
    add_heading_with_color(doc, 'III. 5 nghiệp vụ HR cần cấu hình', level=3, color=(204, 102, 0))
    
    khải_usecases = [
        ('1. Tính tổng quỹ lương theo phòng ban', 'Modeling + Relationship: Tạo calculated field tổng lương, Relationship: Employee_Profile → Department'),
        ('2. Nhân viên > 10 năm chưa thăng chức', 'SQL Pair + Instruction: years_at_company > 10 AND years_since_last_promotion > 5'),
        ('3. So sánh tỷ lệ nghỉ việc: làm thêm giờ vs không', 'SQL Pair: GROUP BY OverTime, HAVING Attrition prediction'),
        ('4. Nhân viên lương bất thường so với phòng ban', 'Instruction: "Bất thường" = chênh lệch > 1.5 độ lệch chuẩn'),
        ('5. Báo cáo tổng hợp (tổng NV, số nghỉ việc dự báo, tỷ lệ churn)', 'SQL Pair tổng hợp: Multi-table aggregation')
    ]
    
    for title, desc in khải_usecases:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(f': {desc}')
    
    add_heading_with_color(doc, 'IV. Source code bắt buộc đọc', level=3, color=(204, 102, 0))
    
    khải_source = [
        'WrenAI/docker/docker-compose.yaml (~120 dòng)',
        'WrenAI/docker/.env (các biến môi trường)',
        'WrenAI/docker/config.yaml (cấu hình 29 pipelines)',
        'WrenAI/wren-ui/src/apollo/server/adaptors/ibisAdaptor.ts (658 dòng)',
        'WrenAI/wren-ai-service/src/providers/engine/wren.py (351 dòng)',
        'WrenAI/wren-ai-service/src/__main__.py (101 dòng)',
        'WrenAI/wren-ai-service/src/globals.py (341 dòng)',
        'WrenAI/docker/bootstrap/init.sh'
    ]
    
    for source in khải_source:
        doc.add_paragraph(source, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== HÂN =====
    add_heading_with_color(doc, '5.2. Hân (K234060691) — Semantic Layer & Vector Store Specialist', level=2)
    
    add_heading_with_color(doc, 'I. Công việc chính', level=3, color=(204, 102, 0))
    hân_jobs = [
        'Chạy lại toàn bộ dự án, đảm bảo deploy thành công (Qdrant index db_schema = 52 documents).',
        'Tạo 5 Relationships phức tạp trong Wren UI (VD: Self-join, Multi-hop join, Calculated field).',
        'Tạo 5 Calculated Fields mới (VD: Age_Group, Salary_Range, Tenure_Category).',
        'Kiểm tra vector hóa Description vào Qdrant bằng cách gọi API localhost:6333.',
        'Định nghĩa 5 nghiệp vụ HR mới và demo thành công.'
    ]
    for job in hân_jobs:
        doc.add_paragraph(job, style='List Number')
    
    add_heading_with_color(doc, 'II. Deep dive questions (trả lời trong Docx)', level=3, color=(204, 102, 0))
    
    hân_questions = [
        'Semantic Layer giải quyết bài toán gì cho Text-to-SQL mà Raw Schema không làm được?',
        'MDL (Model Definition Language) gồm những concepts nào (model, column, relationship, metric, view)?',
        'File db_schema.py trong pipelines/indexing thực hiện vector hóa schema bằng cách nào? Mỗi document chứa những gì?',
        'File db_schema_retrieval.py thực hiện retrieval 2 pha như thế nào (Table retrieval + Column selection)?',
        'Qdrant lưu trữ những collection nào? Mỗi collection có bao nhiêu documents?',
        'Cosine similarity được sử dụng như thế nào để tìm bảng liên quan nhất?',
        'Làm sao AI biết "Attrition = Yes" nghĩa là "Nghỉ việc"? Vai trò của Description và Alias?',
        'Khi nào cần recreate_index = true và khi nào đặt false? Ảnh hưởng thế nào đến thời gian deploy?',
        'Context nào cần cung cấp cho AI để nó hiểu quy trình nghiệp vụ HR?'
    ]
    
    for idx, q in enumerate(hân_questions, 1):
        doc.add_paragraph(f'Q{idx}: {q}', style='List Bullet')
    
    add_heading_with_color(doc, 'III. 5 nghiệp vụ HR cần cấu hình', level=3, color=(204, 102, 0))
    
    hân_usecases = [
        ('1. Phân nhóm nhân viên theo độ tuổi', 'Calculated field: Age_Group = CASE WHEN age < 30 THEN "Young" ...'),
        ('2. So sánh lương với trung bình phòng ban', 'Relationship + Calculated field: Salary_vs_DeptAvg'),
        ('3. Nhân viên undervalued (kinh nghiệm cao, level thấp)', 'Instruction: total_working_years > 10 AND job_level <= 2'),
        ('4. Work-life balance vs attrition theo phòng ban', 'SQL Pair: Multi-table join + GROUP BY'),
        ('5. High Performer at Risk', 'SQL Pair: performance_rating >= 3 AND risk_level IN ("High","Critical")')
    ]
    
    for title, desc in hân_usecases:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(f': {desc}')
    
    add_heading_with_color(doc, 'IV. Source code bắt buộc đọc', level=3, color=(204, 102, 0))
    
    hân_source = [
        'WrenAI/wren-ai-service/src/pipelines/indexing/db_schema.py (393 dòng)',
        'WrenAI/wren-ai-service/src/pipelines/retrieval/db_schema_retrieval.py (520 dòng)',
        'WrenAI/wren-ai-service/src/providers/document_store/qdrant.py (441 dòng)',
        'WrenAI/wren-ai-service/src/providers/embedder/litellm.py (202 dòng)',
        'WrenAI/wren-mdl/mdl.schema.json (472 dòng)',
        'WrenAI/wren-ui/src/utils/modelingHelper.ts (80 dòng)',
        'WrenAI/wren-ui/src/pages/modeling.tsx',
        'WrenAI/wren-ui/src/pages/setup/relationships.tsx'
    ]
    
    for source in hân_source:
        doc.add_paragraph(source, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== NINH =====
    add_heading_with_color(doc, '5.3. Ninh (K234060716) — Agentic Layer & Knowledge Engineer', level=2)
    
    add_heading_with_color(doc, 'I. Công việc chính', level=3, color=(204, 102, 0))
    ninh_jobs = [
        'Chạy lại toàn bộ dự án, kiểm tra pipeline Text-to-SQL hoạt động đúng.',
        'Thêm 5 SQL Pairs mới (dạy AI các câu hỏi khó/lắt léo của HR).',
        'Thêm 5 Instructions mới (quy tắc nghiệp vụ: VD "Luôn lọc nhân viên Active").',
        'Tinh chỉnh config.yaml (temperature, max_tokens) để câu trả lời ổn định.',
        'Demo: hỏi câu hỏi mơ hồ và hệ thống phải tự dùng Instruction để định nghĩa và query đúng.'
    ]
    for job in ninh_jobs:
        doc.add_paragraph(job, style='List Number')
    
    add_heading_with_color(doc, 'II. Deep dive questions (trả lời trong Docx)', level=3, color=(204, 102, 0))
    
    ninh_questions = [
        'File sql_generation.py tổ chức Hamilton DAG gồm những node nào (prompt, generate, post_process)?',
        'Prompt gửi sang Gemini được cấu tạo từ những thành phần nào (System + Schema + Few-shot + Instructions + User Question)?',
        'Cơ chế tự sửa lỗi (Self-Correction) trong sql_correction.py hoạt động như thế nào?',
        'Intent Classification phân loại câu hỏi thành những loại nào (4 loại)?',
        'SQL Pairs được quản lý qua API nào? Quy trình từ lúc thêm đến lúc ảnh hưởng kết quả?',
        'Instructions chia thành 2 loại nào (isGlobal: true/false)? Cho ví dụ cụ thể.',
        'LiteLLM đóng vai trò gì? Tại sao dùng prefix "gemini/" thay vì gọi trực tiếp Gemini API?',
        'Chart generation pipeline tạo biểu đồ Vega-Lite như thế nào?',
        'Làm sao bảo mật thông tin nhân viên khi query qua LLM (chỉ gửi Metadata, không gửi Raw Data)?'
    ]
    
    for idx, q in enumerate(ninh_questions, 1):
        doc.add_paragraph(f'Q{idx}: {q}', style='List Bullet')
    
    add_heading_with_color(doc, 'III. 5 nghiệp vụ HR cần cấu hình', level=3, color=(204, 102, 0))
    
    ninh_usecases = [
        ('1. Top 5 nhân viên có nguy cơ nghỉ việc cao nhất', 'SQL Pair: CTE + ROW_NUMBER + risk_level'),
        ('2. Nhân viên có điểm burnout nguy hiểm', 'SQL Pair + Instruction: Công thức burnout = overtime*3 + business_travel*2 + years_since_last_promotion*1.5'),
        ('3. Lọc mặc định theo risk_level High/Critical', 'Instruction (Global): "Khi hỏi về rủi ro, mặc định chỉ hiển thị High hoặc Critical"'),
        ('4. So sánh tỷ lệ nghỉ việc dự báo giữa các job_role', 'SQL Pair: Multi-group aggregation'),
        ('5. Mặc định sử dụng monthly_income cho lương', 'Instruction (Global): "Luôn dùng monthly_income, không dùng daily_rate"')
    ]
    
    for title, desc in ninh_usecases:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(f': {desc}')
    
    add_heading_with_color(doc, 'IV. Source code bắt buộc đọc', level=3, color=(204, 102, 0))
    
    ninh_source = [
        'WrenAI/wren-ai-service/src/pipelines/generation/sql_generation.py (234 dòng)',
        'WrenAI/wren-ai-service/src/pipelines/generation/sql_correction.py (201 dòng)',
        'WrenAI/wren-ai-service/src/pipelines/generation/intent_classification.py (401 dòng)',
        'WrenAI/wren-ai-service/src/pipelines/generation/chart_generation.py',
        'WrenAI/wren-ai-service/src/pipelines/generation/sql_answer.py',
        'WrenAI/wren-ai-service/src/providers/llm/litellm.py (167 dòng)',
        'WrenAI/docker/config.yaml (cấu hình temperature, max_tokens)',
        'WrenAI/wren-ai-service/src/config.py (122 dòng)',
        'WrenAI/wren-ai-service/src/web/v1/routers/ask.py (80 dòng)',
        'WrenAI/wren-ui/src/pages/knowledge/question-sql-pairs.tsx',
        'WrenAI/wren-ui/src/pages/knowledge/instructions.tsx'
    ]
    
    for source in ninh_source:
        doc.add_paragraph(source, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== GIA =====
    add_heading_with_color(doc, '5.4. Gia (K234060689) — Business Insights Analyst', level=2)
    
    add_heading_with_color(doc, 'I. Công việc chính', level=3, color=(204, 102, 0))
    gia_jobs = [
        'Giải thích từng cell trong Notebook (18 cells: input, output, công thức, thư viện)',
        'Trích xuất Business Insights từ Feature Importance và Model Performance',
        'Tạo biểu đồ: Feature Importance (bar), Correlation Heatmap, Risk Distribution',
        'Viết Report phân tích chi tiết HR analytics (Q1-Q9)',
        'Chuẩn bị slide trình bày Business Insights (5-7 slides)'
    ]
    for job in gia_jobs:
        doc.add_paragraph(job, style='List Number')
    
    add_heading_with_color(doc, 'II. Deep dive questions (trả lời trong Docx)', level=3, color=(204, 102, 0))
    
    gia_questions = [
        'Insight từ feature importance giúp gì cho HR Director ra quyết định chiến lược?',
        'Chi phí thay thế 1 nhân sự là bao nhiêu? (Tìm số liệu SHRM, Gallup, Deloitte)',
        'Data Leakage là gì? Tại sao tách Train/Test bình thường lại sai?',
        'OOF (Out-of-Fold) giúp mô phỏng Production như thế nào?',
        'Chỉ số Recall quan trọng hơn Precision không? Tại sao?',
        'SMOTE hoạt động thế nào? Tại sao cần xử lý mất cân bằng dữ liệu?',
        'Random Forest Classifier được chọn vì lý do gì? So sánh với Logistic Regression và XGBoost.',
        'Các hyperparameters (n_estimators=300, max_depth=15, class_weight="balanced") có ý nghĩa gì?',
        'Thang đo rủi ro (Low, Medium, High, Critical) xây dựng dựa trên cơ sở nào?'
    ]
    
    for idx, q in enumerate(gia_questions, 1):
        doc.add_paragraph(f'Q{idx}: {q}', style='List Bullet')
    
    add_heading_with_color(doc, 'III. Source code bắt buộc đọc', level=3, color=(204, 102, 0))
    
    gia_source = [
        'notebooks/HR_Analytics_Project_Final.ipynb (18 cells)',
        'notebooks/WA_Fn-UseC_-HR-Employee-Attrition.csv (35 features, 1470 rows)',
        'legacy/init-db.sql (178 dòng — bảng hr_training_data)',
        'legacy/create_actionable_views.sql (33 dòng — VIEW v_employee_actionable_insights)'
    ]
    
    for source in gia_source:
        doc.add_paragraph(source, style='List Bullet')
    
    add_heading_with_color(doc, 'IV. Deliverables', level=3, color=(204, 102, 0))
    
    gia_deliver = [
        'Docx Report trả lời Q1-Q9 (có biểu đồ, số liệu, dẫn chứng)',
        'Giải thích từng cell notebook (cell 1-18)',
        'Biểu đồ feature importance (bar chart), correlation heatmap',
        'Bảng phân tích Risk theo Department',
        'Slide trình bày Business Insights (5-7 slides)'
    ]
    
    for deliver in gia_deliver:
        doc.add_paragraph(deliver, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== UYÊN =====
    add_heading_with_color(doc, '5.5. Uyên (K234060737) — MLOps Engineer & Project Architect', level=2)
    
    add_heading_with_color(doc, 'I. Công việc chính', level=3, color=(204, 102, 0))
    uyên_jobs = [
        'Vẽ sơ đồ Cây Dự Án (Project Anatomy) — 3 layers',
        'Vẽ sơ đồ Data Flow end-to-end (từ SQL Server → AI Pipeline → UI)',
        'Vẽ sơ đồ MLOps Workflow (monitoring, retraining, data validation)',
        'Viết Docx tổng quan giải thích architecture cốt lõi',
        'Chuẩn bị slide trình bày tổng quan dự án (7-10 slides)'
    ]
    for job in uyên_jobs:
        doc.add_paragraph(job, style='List Number')
    
    add_heading_with_color(doc, 'II. Deep dive questions (trả lời trong Docx)', level=3, color=(204, 102, 0))
    
    uyên_questions = [
        'Mục tiêu cốt lõi của dự án? Chuyển đổi HR Thụ động → Chủ động?',
        'Dân chủ hóa dữ liệu (Data Democratization) nghĩa là gì?',
        'Tại sao bảo mật quan trọng? Chỉ gửi Metadata cho LLM?',
        'Luồng ETL hoạt động cụ thể như thế nào?',
        'Model Drift là gì? Khi nào cần Retrain?',
        'Data Validation chống "Garbage in, Garbage out"?',
        'Trigger logic: Tại sao chạy theo tháng?',
        'Giám sát phân phối dữ liệu (Monitoring)?',
        'Human-in-the-loop: AI hỗ trợ hay thay thế HR Director?',
        'Production Tools: Notebook → Apache Airflow / SQL Server Agent Job?',
        'Tại sao nên dùng Local LLM (Ollama) thay vì Cloud API?',
        'Giá trị Email Insight và HTML Report?'
    ]
    
    for idx, q in enumerate(uyên_questions, 1):
        doc.add_paragraph(f'Q{idx}: {q}', style='List Bullet')
    
    add_heading_with_color(doc, 'III. Source code bắt buộc đọc', level=3, color=(204, 102, 0))
    
    uyên_source = [
        'Toàn bộ cấu trúc thư mục repo (folder structure)',
        'legacy/init-db.sql (178 dòng)',
        'legacy/create_actionable_views.sql (33 dòng)',
        'legacy/setup_db_mail_template.sql (74 dòng — Database Mail, Gmail SMTP)',
        'notebooks/HR_Analytics_Project_Final.ipynb (luồng ETL + ML)',
        'WrenAI/docker/docker-compose.yaml (6 services overview)',
        'WrenAI/docker/config.yaml (LLM, Embedder, Qdrant)',
        'HR_Analytics.bak (SQL Server backup file)'
    ]
    
    for source in uyên_source:
        doc.add_paragraph(source, style='List Bullet')
    
    add_heading_with_color(doc, 'IV. Deliverables', level=3, color=(204, 102, 0))
    
    uyên_deliver = [
        'Docx trả lời Q1-Q12 (có sơ đồ minh họa)',
        'Sơ đồ Cây Dự Án (Project Anatomy) — 3 layers',
        'Sơ đồ Data Flow end-to-end',
        'Sơ đồ MLOps Workflow',
        'Slide trình bày tổng quan dự án (7-10 slides)',
        'Bảng PCCV hoàn chỉnh (Excel)'
    ]
    
    for deliver in uyên_deliver:
        doc.add_paragraph(deliver, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== THROUGH-BACK CHUNG =====
    add_heading_with_color(doc, '6. Through-back chung (Q&A Showcase) — Khải, Hân, Ninh', level=1)
    
    doc.add_paragraph('Tất cả 3 thành viên (Khải, Hân, Ninh) cùng nghiên cứu và chuẩn bị trả lời các câu hỏi sau trong buổi showcase:')
    doc.add_paragraph()
    
    through_back_qs = [
        'Tại sao chọn Wren AI thay vì LangChain SQL Agent hoặc LlamaIndex? Wren AI có gì khác biệt?',
        'Semantic Layer giải quyết bài toán gì cho Text-to-SQL mà các giải pháp khác không có?',
        'Làm sao để bảo mật thông tin nhân viên khi query qua LLM? Chỉ gửi Metadata hay gửi cả Raw Data?',
        'Các tính năng chính của Wren AI là gì, sắp xếp trong folder nào, chạy class chính nào?',
        'Làm sao kết nối được tới SQL Server qua gì (Ibis Server, Connection String)?',
        'Kiến trúc 4 layer (Data, Semantic, Agentic, Representation) hoạt động ra sao?',
        'Context nào cần cung cấp cho AI để nó hiểu quy trình nghiệp vụ HR?',
        'Hệ thống có thể scale cho dataset lớn hơn (10K+ employees) không? Cần nâng cấp gì?'
    ]
    
    for idx, q in enumerate(through_back_qs, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(q).font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    doc.add_paragraph('Phân công cụ thể:')
    
    assignment_table = doc.add_table(rows=4, cols=2)
    assignment_table.style = 'Light Grid Accent 1'
    add_table_border(assignment_table)
    
    assignment_data = [
        ('Người', 'Câu hỏi'),
        ('Khải', 'Q1, Q4, Q5, Q6 (Kiến trúc + Infrastructure)'),
        ('Hân', 'Q2, Q7 (Semantic Layer + Context)'),
        ('Ninh', 'Q3, Q8 (Bảo mật + Scale)')
    ]
    
    for idx, (person, qs) in enumerate(assignment_data):
        cells = assignment_table.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            cells[0].text = person
            cells[1].text = qs
    
    doc.add_page_break()
    
    # ===== PHẦN 7: HƯỚNG DẪN CẤU HÌNH =====
    add_heading_with_color(doc, '7. Hướng dẫn cấu hình 5 nghiệp vụ HR trên Wren AI UI', level=1)
    
    add_heading_with_color(doc, '7.1. Cách tạo Relationship (cho Khải, Hân)', level=2, color=(0, 102, 204))
    
    rel_steps = [
        'Mở Wren AI UI → tab Modeling',
        'Click vào model cần tạo relationship (VD: Employee_Profile)',
        'Click Add Relationship',
        'Chọn model liên kết (VD: Department)',
        'Chọn Join Type: MANY_TO_ONE',
        'Chọn Join Column: Employee_Profile.department → Department.department_name',
        'Click Save → Click Deploy (góc trên phải) để sync vào AI'
    ]
    
    for idx, step in enumerate(rel_steps, 1):
        doc.add_paragraph(step, style='List Number')
    
    add_heading_with_color(doc, '7.2. Cách tạo Calculated Field (cho Khải, Hân)', level=2, color=(0, 102, 204))
    
    calc_steps = [
        'Mở Modeling → Chọn model (VD: Employee_Profile)',
        'Click Add Calculated Field',
        'Nhập tên: VD Age_Group',
        'Nhập biểu thức SQL: CASE WHEN age < 30 THEN "Young" WHEN age < 45 THEN "Mid" ELSE "Senior" END',
        'Chọn data type: VARCHAR',
        'Click Save → Deploy'
    ]
    
    for idx, step in enumerate(calc_steps, 1):
        doc.add_paragraph(step, style='List Number')
    
    add_heading_with_color(doc, '7.3. Cách thêm SQL Pair (cho Ninh)', level=2, color=(0, 102, 204))
    
    sql_pair_steps = [
        'Mở Knowledge tab → Question SQL Pairs',
        'Click + Add SQL Pair',
        'Nhập Question: "Top 5 nhân viên có nguy cơ nghỉ việc cao nhất tháng này"',
        'Nhập SQL: SELECT TOP 5 ... JOIN ... WHERE risk_level IN ("High", "Critical") ORDER BY probability_score DESC',
        'Click Save → Hệ thống tự động vector hóa và lưu vào Qdrant'
    ]
    
    for idx, step in enumerate(sql_pair_steps, 1):
        doc.add_paragraph(step, style='List Number')
    
    add_heading_with_color(doc, '7.4. Cách thêm Instruction (cho Ninh)', level=2, color=(0, 102, 204))
    
    instr_steps = [
        'Mở Knowledge tab → Instructions',
        'Click + Add Instruction',
        'Nhập Instruction: "Khi hỏi về rủi ro nghỉ việc, mặc định chỉ hiển thị nhân viên có risk_level là High hoặc Critical"',
        'Chọn Type: Global (áp dụng cho mọi câu hỏi)',
        'Click Save → Deploy'
    ]
    
    for idx, step in enumerate(instr_steps, 1):
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # ===== PHẦN 8: HƯỚNG DẪN TEST =====
    add_heading_with_color(doc, '8. Hướng dẫn test demo performance', level=1)
    
    add_heading_with_color(doc, '8.1. Test trên Public Link', level=2, color=(0, 102, 204))
    
    doc.add_paragraph('Mở link: https://certification-lows-spy-tension.trycloudflare.com')
    doc.add_paragraph('Click + New để tạo thread mới')
    doc.add_paragraph('Nhập câu hỏi test cho từng nghiệp vụ:')
    doc.add_paragraph()
    
    add_heading_with_color(doc, '✓ Test Khải (Infrastructure):', level=3, color=(102, 102, 102))
    test_khải = [
        'Tổng quỹ lương của từng phòng ban là bao nhiêu?',
        'Nhân viên nào có thời gian làm việc > 10 năm nhưng chưa được thăng chức?',
        'So sánh tỷ lệ nghỉ việc giữa nhân viên làm thêm giờ và không làm thêm giờ'
    ]
    for q in test_khải:
        doc.add_paragraph(q, style='List Bullet')
    
    add_heading_with_color(doc, '✓ Test Hân (Semantic Layer):', level=3, color=(102, 102, 102))
    test_hân = [
        'Phân nhóm nhân viên theo độ tuổi và tỷ lệ nghỉ việc của từng nhóm?',
        'Nhân viên nào có lương thấp hơn trung bình phòng ban?',
        'Danh sách High Performer at Risk?'
    ]
    for q in test_hân:
        doc.add_paragraph(q, style='List Bullet')
    
    add_heading_with_color(doc, '✓ Test Ninh (Agentic Layer):', level=3, color=(102, 102, 102))
    test_ninh = [
        'Top 5 nhân viên có nguy cơ nghỉ việc cao nhất?',
        'Nhân viên nào có điểm burnout nguy hiểm?',
        'So sánh tỷ lệ nghỉ việc dự báo giữa các job role?'
    ]
    for q in test_ninh:
        doc.add_paragraph(q, style='List Bullet')
    
    add_heading_with_color(doc, '8.2. Kiểm tra performance', level=2, color=(0, 102, 204))
    
    perf_checks = [
        'View SQL tab: SQL sinh ra có đúng logic không?',
        'Chart tab: Biểu đồ có hiển thị đúng dữ liệu không?',
        'Answer tab: Câu trả lời có chính xác, có insights không?',
        'Answer preparation steps: Kiểm tra AI đi qua bao nhiêu bước (thường 3 steps)'
    ]
    
    for check in perf_checks:
        doc.add_paragraph(check, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== PHẦN 9: DELIVERABLES =====
    add_heading_with_color(doc, '9. Yêu cầu đầu ra (Deliverables)', level=1)
    
    add_heading_with_color(doc, '9.1. Khải, Hân, Ninh (AI Engineering)', level=2, color=(0, 102, 204))
    
    deliver_table_1 = doc.add_table(rows=6, cols=2)
    deliver_table_1.style = 'Light Grid Accent 1'
    add_table_border(deliver_table_1)
    
    deliver_1_data = [
        ('Deliverable', 'Mô tả'),
        ('Docx Deep Dive', 'Trả lời tất cả deep dive questions (Q1-Q8/Q9), dẫn chứng source code, số dòng'),
        ('5 nghiệp vụ HR', 'Cấu hình trên Wren AI UI, chụp screenshot kết quả'),
        ('Sơ đồ kỹ thuật', 'Kiến trúc, pipeline, indexing flow (draw.io / Mermaid)'),
        ('Demo live', 'Hỏi câu hỏi trên public link, show kết quả'),
        ('Through-back Q&A', 'Chuẩn bị trả lời 8 câu hỏi vấn đáp showcase')
    ]
    
    for idx, (deliver, desc) in enumerate(deliver_1_data):
        cells = deliver_table_1.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            cells[0].text = deliver
            cells[1].text = desc
    
    add_heading_with_color(doc, '9.2. Gia (Business Analytics)', level=2, color=(0, 102, 204))
    
    deliver_table_2 = doc.add_table(rows=6, cols=2)
    deliver_table_2.style = 'Light Grid Accent 1'
    add_table_border(deliver_table_2)
    
    deliver_2_data = [
        ('Deliverable', 'Mô tả'),
        ('Docx Report', 'Trả lời Q1-Q9, biểu đồ, số liệu, dẫn chứng'),
        ('Giải thích Notebook', 'Cell 1-18: mục đích, input, output, thư viện'),
        ('Biểu đồ', 'Feature importance, correlation heatmap, Risk by Department'),
        ('Slide', 'Business Insights (5-7 slides)'),
        ('', '')
    ]
    
    for idx, (deliver, desc) in enumerate(deliver_2_data):
        cells = deliver_table_2.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            cells[0].text = deliver
            cells[1].text = desc
    
    add_heading_with_color(doc, '9.3. Uyên (MLOps & Architecture)', level=2, color=(0, 102, 204))
    
    deliver_table_3 = doc.add_table(rows=7, cols=2)
    deliver_table_3.style = 'Light Grid Accent 1'
    add_table_border(deliver_table_3)
    
    deliver_3_data = [
        ('Deliverable', 'Mô tả'),
        ('Docx Tổng quan', 'Trả lời Q1-Q12, dẫn chứng file repo, giải thích architecture'),
        ('3 Sơ đồ', 'Cây Dự Án, Data Flow, MLOps Workflow'),
        ('Slide', 'Tổng quan dự án (7-10 slides)'),
        ('Bảng PCCV', 'File Excel PCCV_HR_Analytics_v3.xlsx hoàn chỉnh'),
        ('', ''),
        ('', '')
    ]
    
    for idx, (deliver, desc) in enumerate(deliver_3_data):
        cells = deliver_table_3.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            cells[0].text = deliver
            cells[1].text = desc
    
    doc.add_page_break()
    
    # ===== PHẦN 10: TIMELINE =====
    add_heading_with_color(doc, '10. Timeline, Checklist & Submission', level=1)
    
    add_heading_with_color(doc, '10.1. Timeline chi tiết (11/02 - 16/02/2026)', level=2, color=(0, 102, 204))
    
    timeline_table = doc.add_table(rows=7, cols=4)
    timeline_table.style = 'Light Grid Accent 1'
    add_table_border(timeline_table)
    
    timeline_data = [
        ('Ngày', 'Hoạt động', 'Người phụ trách', 'Checklist'),
        ('11/02 (T3)', 'Nhận tài liệu, clone repo, chạy dự án', 'Tất cả', '☐ Repo chạy thành công'),
        ('12/02 (T4)', 'Đọc source code theo phân công', 'Tất cả', '☐ Ghi chú cá nhân'),
        ('13/02 (T5)', 'Cấu hình 5 nghiệp vụ HR, test trên UI', 'Khải, Hân, Ninh', '☐ 5 nghiệp vụ active'),
        ('13/02 (T5)', 'Giải thích notebook, tạo biểu đồ', 'Gia', '☐ Draft report + biểu đồ'),
        ('13/02 (T5)', 'Vẽ sơ đồ, viết phần tổng quan', 'Uyên', '☐ Draft 3 sơ đồ'),
        ('14/02 (T6)', 'Viết Docx deep dive, Through-back prep', 'Tất cả', '☐ Draft Docx hoàn chỉnh')
    ]
    
    for idx, row_data in enumerate(timeline_data):
        cells = timeline_table.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)
        else:
            for cell_idx, value in enumerate(row_data):
                cells[cell_idx].text = value
    
    # Add more timeline rows
    timeline_table.add_row()
    row = timeline_table.rows[-1]
    row.cells[0].text = '15/02 (T7)'
    row.cells[1].text = 'Hoàn thiện Docx, tạo Slide, review chéo'
    row.cells[2].text = 'Tất cả'
    row.cells[3].text = '☐ Docx + Slide final'
    
    timeline_table.add_row()
    row = timeline_table.rows[-1]
    row.cells[0].text = '16/02 (CN)'
    row.cells[1].text = 'Push deliverables lên GitLab, dry run showcase'
    row.cells[2].text = 'Tất cả'
    row.cells[3].text = '☐ Tất cả trên GitLab'
    
    add_heading_with_color(doc, '10.2. Submission Format', level=2, color=(0, 102, 204))
    
    doc.add_paragraph('Mỗi thành viên tạo 1 folder với tên {Tên_người} chứa:')
    doc.add_paragraph()
    
    submission_items = [
        'Docx Deep Dive (Q&A + trích dẫn source code + biểu đồ + screenshot)',
        '5 nghiệp vụ HR screenshots',
        'Sơ đồ kỹ thuật (draw.io / Mermaid / PNG)',
        'Slide PowerPoint (nếu có)',
        'Source code notes (các file source code được đọc, ghi chú quan trọng)'
    ]
    
    for item in submission_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Push tất cả lên GitLab branch: hr_domain_research bằng lệnh:')
    p = doc.add_paragraph()
    p.add_run('git add . && git commit -m "Deliverables: [Tên_người] Q&A + Use Cases" && git push').font.name = 'Courier New'
    
    add_heading_with_color(doc, '10.3. Evaluation Criteria', level=2, color=(0, 102, 204))
    
    evaluation_criteria = [
        ('Hoàn thiện (Completeness)', '100%: Tất cả Q&A đã trả lời, 5 use cases cấu hình thành công'),
        ('Chính xác (Accuracy)', '100%: Deep dive questions trả lời đúng, dẫn chứng source code chính xác'),
        ('Insight (Insights)', '100%: Giải thích tại sao, không chỉ cơ chế, kết nối với bài toán HR'),
        ('Trình bày (Presentation)', '100%: Docx rõ ràng, sơ đồ đẹp, slide chuyên nghiệp'),
        ('Collaboration', '100%: Through-back chung được chuẩn bị kỹ lưỡng')
    ]
    
    eval_table = doc.add_table(rows=6, cols=2)
    eval_table.style = 'Light Grid Accent 1'
    add_table_border(eval_table)
    
    for idx, (criterion, desc) in enumerate(evaluation_criteria):
        cells = eval_table.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            cells[0].text = criterion
            cells[1].text = desc
    
    doc.add_page_break()
    
    # ===== FOOTER =====
    add_heading_with_color(doc, 'LIÊN HỆ & HỖ TRỢ', level=1)
    
    contact_info = [
        ('GitLab Repository', 'https://gitlab.com/boygia757-netizen/hr-ai-project'),
        ('Branch', 'hr_domain_research'),
        ('Public Demo Link', 'https://certification-lows-spy-tension.trycloudflare.com'),
        ('Local UI URL', 'http://localhost:3000'),
        ('AI Service Health', 'http://localhost:5555/health'),
        ('Qdrant Dashboard', 'http://localhost:6333/dashboard'),
        ('Hướng dẫn bổ sung', 'HUONG_DAN_TEAM_DEV.md (trong repo)')
    ]
    
    for label, value in contact_info:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        if 'http' in value:
            run = p.add_run(value)
            run.font.color.rgb = RGBColor(0, 0, 255)
        else:
            p.add_run(value)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph('Tài liệu này được tạo tự động ngày ' + datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S'))
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    return doc

if __name__ == '__main__':
    doc = create_pccv_doc()
    output_path = 'HUONG_DAN_TEAM_DEV.docx'
    doc.save(output_path)
    print(f'✓ Tài liệu đã tạo thành công: {output_path}')
    print(f'✓ Số trang: ~30 pages')
    print(f'✓ Nội dung: Bảng PCCV, Phân công, Deep Dive, Through-back, Timeline')
