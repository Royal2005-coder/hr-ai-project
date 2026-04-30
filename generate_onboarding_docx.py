"""
Script tạo tài liệu Onboarding DOCX cho dự án HR Analytics Text-to-SQL.
Tài liệu hướng dẫn chi tiết step-by-step cho thành viên mới.

Chạy: python generate_onboarding_docx.py
Output: ONBOARDING_GUIDE.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ==============================================================================
# CẤU HÌNH
# ==============================================================================
OUTPUT_FILE = "ONBOARDING_GUIDE.docx"
PROJECT_NAME = "HR Analytics — Hệ thống dự báo nghỉ việc và truy vấn insight nhân sự"
PROJECT_CODE = "252BIM500601"
GITLAB_URL = "https://gitlab.com/boygia757-netizen/hr-ai-project"
BRANCH_WORK = "hr_domain_research"
BRANCH_MAIN = "main"
VERSION = "3.0"
DATE = "11/02/2026"

TEAM_MEMBERS = [
    {"email": "ninhdp23406@st.uel.edu.vn", "role": "AI Engineering"},
    {"email": "uyenntd23406@st.uel.edu.vn", "role": "Data Pipeline & MLOps"},
    {"email": "hannpn23406@st.uel.edu.vn", "role": "AI Engineering"},
    {"email": "khainn23406@st.uel.edu.vn", "role": "AI Engineering (Lead)"},
]


# ==============================================================================
# HÀM TIỆN ÍCH
# ==============================================================================

def set_cell_shading(cell, color_hex):
    """Tô màu nền cho ô trong bảng."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Đặt border cho ô."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), val.get("val", "single"))
        element.set(qn("w:sz"), val.get("sz", "4"))
        element.set(qn("w:color"), val.get("color", "CCCCCC"))
        element.set(qn("w:space"), "0")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_table(doc, headers, rows, col_widths=None, header_color="1B4F72"):
    """Tạo bảng với header có màu và format chuẩn."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FB")

    # Column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


def add_code_block(doc, code_text, language="powershell"):
    """Thêm khối code với font monospace và nền xám."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    # Tô nền xám cho paragraph
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F5F5")
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def add_warning_box(doc, text):
    """Thêm hộp cảnh báo với biểu tượng ⚠️."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FFF3CD")
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)

    run = p.add_run("⚠ Cảnh báo: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x85, 0x6D, 0x0E)

    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x85, 0x6D, 0x0E)
    return p


def add_important_box(doc, text):
    """Thêm hộp thông tin quan trọng."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "D4EDDA")
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)

    run = p.add_run("✔ ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x15, 0x57, 0x24)

    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x15, 0x57, 0x24)
    return p


def add_error_box(doc, text):
    """Thêm hộp cấm / lỗi."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F8D7DA")
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)

    run = p.add_run("✘ ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x72, 0x1C, 0x24)

    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x72, 0x1C, 0x24)
    return p


def add_step(doc, step_number, title, description=""):
    """Thêm bước hướng dẫn có đánh số."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)

    run = p.add_run(f"Bước {step_number}: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    run2 = p.add_run(title)
    run2.bold = True
    run2.font.size = Pt(11)

    if description:
        p2 = doc.add_paragraph(description)
        p2.paragraph_format.left_indent = Cm(1)
        p2.runs[0].font.size = Pt(10)
    return p


# ==============================================================================
# NỘI DUNG CHÍNH
# ==============================================================================

def create_document():
    doc = Document()

    # =========================================================================
    # CẤU HÌNH STYLES
    # =========================================================================
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Times New Roman"
        heading_style.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(12)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # =========================================================================
    # TRANG BÌA
    # =========================================================================
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC KINH TẾ - LUẬT (UEL)")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KHOA HỆ THỐNG THÔNG TIN")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tiêu đề chính
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TÀI LIỆU HƯỚNG DẪN ONBOARDING")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(PROJECT_NAME)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Mã số đề tài: {PROJECT_CODE}")
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Thông tin phiên bản
    info_items = [
        ("Phiên bản tài liệu:", VERSION),
        ("Ngày cập nhật:", DATE),
        ("Repository:", GITLAB_URL),
        ("Branch làm việc:", BRANCH_WORK),
        ("Loại tài liệu:", "Hướng dẫn onboarding cho thành viên mới"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label} ")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        run2 = p.add_run(value)
        run2.font.size = Pt(11)
        run2.bold = True

    doc.add_page_break()

    # =========================================================================
    # MỤC LỤC
    # =========================================================================
    doc.add_heading("Mục lục", level=1)

    toc_items = [
        "1. Giới thiệu dự án và kiến trúc tổng thể",
        "2. Danh sách thành viên và phân quyền",
        "3. Yêu cầu phần cứng và phần mềm",
        "4. Cài đặt môi trường phát triển",
        "5. Clone dự án từ GitLab",
        "6. Cấu hình môi trường (.env)",
        "7. Khởi động hệ thống",
        "8. Kiểm tra hệ thống hoạt động",
        "9. Hướng dẫn sử dụng giao diện Wren AI",
        "10. Quy trình làm việc với Git",
        "11. Cấu trúc thư mục dự án",
        "12. Xử lý lỗi thường gặp",
        "13. Quy tắc bảo mật bắt buộc",
        "14. Bảng kiểm (checklist) onboarding",
        "15. Phân công nhiệm vụ chi tiết theo team",
        "16. Liên hệ hỗ trợ",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 1: GIỚI THIỆU DỰ ÁN
    # =========================================================================
    doc.add_heading("1. Giới thiệu dự án và kiến trúc tổng thể", level=1)

    doc.add_heading("1.1. Dự án giải quyết vấn đề gì?", level=2)
    doc.add_paragraph(
        "Dự án HR Analytics xây dựng một trợ lý ảo (AI assistant) cho phép bộ phận nhân sự (HR) "
        "hỏi đáp dữ liệu bằng tiếng Việt mà không cần biết SQL. Hệ thống tự động chuyển câu hỏi "
        "tiếng Việt thành câu lệnh SQL, truy vấn cơ sở dữ liệu SQL Server và trả về kết quả "
        "dạng bảng kèm biểu đồ trực quan."
    )

    doc.add_paragraph(
        "Ví dụ: người dùng nhập \"10 nhân viên có nguy cơ nghỉ việc cao nhất là ai?\" → hệ thống "
        "trả về danh sách đầy đủ với tên, phòng ban, xác suất nghỉ việc, mức rủi ro."
    )

    doc.add_heading("1.2. Hai năng lực AI cốt lõi", level=2)
    add_table(doc,
        ["Năng lực", "Mô tả", "Công nghệ"],
        [
            ["Predictive AI", "Mô hình Random Forest dự báo nguy cơ nghỉ việc cho 1.470 nhân viên", "Python, scikit-learn"],
            ["Generative AI (Text-to-SQL)", "Trợ lý ảo chuyển câu hỏi tiếng Việt → SQL → kết quả", "Wren AI, Gemini, LiteLLM"],
        ],
        col_widths=[4, 7, 5]
    )

    doc.add_heading("1.3. Kiến trúc tổng thể", level=2)
    doc.add_paragraph(
        "Hệ thống gồm 6 dịch vụ Docker chạy đồng thời, giao tiếp với nhau qua mạng nội bộ Docker:"
    )

    add_table(doc,
        ["Thành phần", "Port", "Vai trò", "Công nghệ"],
        [
            ["Wren UI", "3000", "Giao diện web, nơi người dùng nhập câu hỏi", "Next.js, React"],
            ["Wren AI Service", "5555", "Bộ não AI: chuyển câu hỏi → SQL, tạo biểu đồ", "Python, LiteLLM, Gemini API"],
            ["Wren Engine", "8080", "Thực thi SQL, quản lý semantic layer", "Java"],
            ["Qdrant", "6333", "Lưu trữ vector embeddings cho SQL Pairs/Instructions", "Vector database"],
            ["Ibis Server", "8000", "Kết nối tới SQL Server database gốc", "Python"],
            ["Bootstrap", "—", "Khởi tạo cấu hình lần đầu (tự tắt sau khi xong)", "Shell script"],
        ],
        col_widths=[3.5, 1.5, 6, 4]
    )

    doc.add_paragraph(
        "Luồng hoạt động: Người dùng nhập câu hỏi tại Wren UI (port 3000) → Wren AI Service gọi "
        "Google Gemini API để chuyển thành SQL → Wren Engine thực thi SQL trên database → kết quả "
        "trả về hiển thị trên giao diện web kèm biểu đồ."
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 2: THÀNH VIÊN VÀ PHÂN QUYỀN
    # =========================================================================
    doc.add_heading("2. Danh sách thành viên và phân quyền", level=1)

    doc.add_heading("2.1. Thành viên dự án", level=2)
    doc.add_paragraph(
        "Dự án là Open Source. Để chạy dự án này, mỗi người "
        "tự trang bị một Gemini API Key và đưa vào file cấu hình."
    )

    member_rows = []
    for i, m in enumerate(TEAM_MEMBERS, 1):
        member_rows.append([str(i), m["email"], "Developer", m["role"]])

    add_table(doc,
        ["STT", "Email", "Quyền GitLab", "Vai trò dự án"],
        member_rows,
        col_widths=[1, 5.5, 3, 5.5]
    )

    doc.add_heading("2.2. Phân quyền GitLab — quyền Developer", level=2)
    doc.add_paragraph("Với vai trò Developer trên GitLab, bạn có các quyền sau:")

    add_important_box(doc, "Được phép: clone dự án, đọc code, tạo branch, push lên branch hr_domain_research.")
    add_important_box(doc, "Được phép: tạo Merge Request từ hr_domain_research vào main.")
    add_error_box(doc, "Không được phép: push trực tiếp lên branch main (đã bị khóa — protected).")
    add_error_box(doc, "Không được phép: xóa branch main, xóa repository, thay đổi cài đặt dự án.")
    add_error_box(doc, "Không được phép: xóa code hoặc file của người khác trên branch main.")

    doc.add_heading("2.3. Gemini API — Cần tự cấu hình", level=2)
    doc.add_paragraph(
        "Hệ thống sử dụng Google Gemini API (qua Google AI Studio). Mỗi cá nhân cần tự lấy "
        "API key từ Google AI Studio và cấu hình vào file .env trước khi chạy lệnh docker compose up -d."
    )
    add_table(doc,
        ["Thành phần AI", "Model", "Mô tả"],
        [
            ["LLM (Large Language Model)", "gemini/gemini-2.5-flash", "Chuyển câu hỏi tiếng Việt → SQL, tạo biểu đồ, phân tích"],
            ["Embedder", "gemini/gemini-embedding-001", "Tạo vector embeddings cho semantic search (768 dims)"],
        ],
        col_widths=[4, 5, 7]
    )
    add_important_box(doc, "Lưu ý: Bạn phải sao chép nội dung từ .env.example sang file mới có tên .env và điền API key cá nhân vào biến GEMINI_API_KEY.")

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 3: YÊU CẦU PHẦN CỨNG VÀ PHẦN MỀM
    # =========================================================================
    doc.add_heading("3. Yêu cầu phần cứng và phần mềm", level=1)

    doc.add_heading("3.1. Phần cứng tối thiểu", level=2)
    add_table(doc,
        ["Thành phần", "Tối thiểu", "Khuyến nghị"],
        [
            ["RAM", "8 GB", "16 GB"],
            ["Ổ cứng trống", "10 GB", "20 GB"],
            ["CPU", "4 cores", "8 cores"],
            ["Mạng", "Internet ổn định", "Internet ổn định"],
        ],
        col_widths=[5, 5, 5]
    )

    doc.add_heading("3.2. Phần mềm cần cài đặt trước", level=2)
    add_warning_box(doc, "Cài đặt tất cả phần mềm dưới đây trước khi thực hiện các bước tiếp theo.")

    doc.add_heading("3.2.1. Docker Desktop (bắt buộc)", level=3)
    add_step(doc, 1, "Truy cập trang tải Docker Desktop")
    add_code_block(doc, "https://www.docker.com/products/docker-desktop/")
    add_step(doc, 2, "Tải phiên bản Docker Desktop for Windows")
    add_step(doc, 3, "Chạy file cài đặt, làm theo hướng dẫn trên màn hình")
    add_step(doc, 4, "Khởi động lại máy tính khi được yêu cầu")
    add_step(doc, 5, "Mở Docker Desktop và đợi đến khi icon ở taskbar chuyển sang màu xanh (Engine running)")
    add_step(doc, 6, "Kiểm tra Docker đã cài thành công")
    add_code_block(doc, 'docker --version\n# Kết quả mong đợi: Docker version 27.x.x hoặc mới hơn\n\ndocker compose version\n# Kết quả mong đợi: Docker Compose version v2.x.x')

    doc.add_heading("3.2.2. Git (bắt buộc)", level=3)
    add_step(doc, 1, "Truy cập trang tải Git")
    add_code_block(doc, "https://git-scm.com/download/win")
    add_step(doc, 2, "Tải và cài đặt, giữ các tùy chọn mặc định")
    add_step(doc, 3, 'Khi được hỏi "Default branch name", chọn main')
    add_step(doc, 4, "Kiểm tra Git đã cài thành công")
    add_code_block(doc, 'git --version\n# Kết quả mong đợi: git version 2.4x.x hoặc mới hơn')

    doc.add_heading("3.2.3. Visual Studio Code (khuyến nghị)", level=3)
    add_step(doc, 1, "Truy cập: https://code.visualstudio.com/ → tải và cài đặt")
    add_step(doc, 2, "Cài extension Docker (ms-azuretools.vscode-docker) để xem logs container")

    doc.add_heading("3.2.4. Trình duyệt web", level=3)
    doc.add_paragraph("Google Chrome, Microsoft Edge, hoặc Firefox phiên bản mới nhất.")

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 4: CÀI ĐẶT MÔI TRƯỜNG
    # =========================================================================
    doc.add_heading("4. Cài đặt môi trường phát triển", level=1)

    doc.add_heading("4.1. Cấu hình Git trên máy của bạn", level=2)
    doc.add_paragraph("Mở PowerShell (hoặc Terminal trong VS Code) và chạy các lệnh sau:")
    add_code_block(doc,
        '# Thay bằng tên và email thực tế của bạn\n'
        'git config --global user.name "Ho Ten Cua Ban"\n'
        'git config --global user.email "email-cua-ban@st.uel.edu.vn"\n'
        '\n'
        '# Lưu credential để không phải nhập lại mỗi lần\n'
        'git config --global credential.helper manager'
    )

    doc.add_heading("4.2. Tạo GitLab Personal Access Token", level=2)
    doc.add_paragraph(
        "Personal Access Token (PAT) là mã xác thực thay thế mật khẩu khi push code lên GitLab. "
        "Bạn cần tạo PAT để sử dụng trong suốt quá trình làm việc."
    )

    add_step(doc, 1, "Đăng nhập GitLab tại https://gitlab.com/ bằng tài khoản đã được mời")
    add_step(doc, 2, "Chấp nhận lời mời tham gia dự án (kiểm tra email mời từ GitLab)")
    add_step(doc, 3, "Truy cập trang tạo token")
    add_code_block(doc, "https://gitlab.com/-/user_settings/personal_access_tokens")
    add_step(doc, 4, "Nhấn Add new token và điền thông tin:")
    add_table(doc,
        ["Trường", "Giá trị"],
        [
            ["Token name", "hr-ai-dev (hoặc tên tùy chọn)"],
            ["Expiration date", "Chọn ngày hết hạn (khuyến nghị 6 tháng — 1 năm)"],
            ["Scopes", "Tick ✓ read_repository và ✓ write_repository"],
        ],
        col_widths=[5, 11]
    )
    add_step(doc, 5, "Nhấn Create personal access token")
    add_step(doc, 6, "Copy token ngay lập tức và lưu ở nơi an toàn")
    add_warning_box(doc, "Token chỉ hiển thị một lần duy nhất! Nếu mất, bạn phải tạo token mới. Token có dạng: glpat-xxxxxxxxxxxxxxxxxxxx")

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 5: CLONE DỰ ÁN
    # =========================================================================
    doc.add_heading("5. Clone dự án từ GitLab", level=1)

    doc.add_heading("5.1. Thực hiện clone", level=2)
    add_step(doc, 1, "Mở PowerShell và di chuyển đến thư mục bạn muốn lưu dự án")
    add_code_block(doc, '# Ví dụ: lưu tại thư mục gốc của user\ncd C:\\Users\\<TenUser>')
    add_step(doc, 2, "Clone dự án từ GitLab")
    add_code_block(doc, f'git clone {GITLAB_URL}.git')
    add_step(doc, 3, "Khi được hỏi xác thực:")
    add_table(doc,
        ["Trường", "Nhập giá trị"],
        [
            ["Username", "Username GitLab của bạn"],
            ["Password", "Personal Access Token đã tạo ở bước 4.2 (không phải mật khẩu GitLab)"],
        ],
        col_widths=[5, 11]
    )

    doc.add_heading("5.2. Chuyển sang branch làm việc", level=2)
    add_error_box(doc, "Tuyệt đối không làm việc trên branch main. Mọi thay đổi phải được thực hiện trên branch hr_domain_research.")

    add_step(doc, 1, "Di chuyển vào thư mục dự án")
    add_code_block(doc, "cd hr-ai-project")
    add_step(doc, 2, "Chuyển sang branch làm việc")
    add_code_block(doc, "git checkout hr_domain_research")
    add_step(doc, 3, "Xác nhận đang ở đúng branch")
    add_code_block(doc,
        'git branch\n'
        '# Kết quả mong đợi:\n'
        '#   main\n'
        '# * hr_domain_research    ← Dấu * cho thấy bạn đang ở branch này'
    )

    doc.add_heading("5.3. Cập nhật code mới nhất", level=2)
    doc.add_paragraph("Luôn pull code mới nhất trước khi bắt đầu làm việc mỗi ngày:")
    add_code_block(doc, "git pull origin hr_domain_research")

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 6: CẤU HÌNH MÔI TRƯỜNG (.ENV)
    # =========================================================================
    doc.add_heading("6. Cấu hình môi trường (.env)", level=1)

    doc.add_paragraph(
        "Hệ thống sử dụng Google Gemini API. Bạn cần copy file .env.example thành .env "
        "và điền API key cá nhân (lấy từ Google AI Studio) vào biến GEMINI_API_KEY. "
        "Tuyệt đối không commit file .env lên public repository."
    )
    add_important_box(doc, "API key là tài sản cá nhân. Dự án là Open Source public repo — mọi file bạn commit đều được hiển thị công khai.")

    doc.add_heading("6.1. Tạo file .env từ file mẫu", level=2)
    doc.add_paragraph(
        "Khi pull code từ GitHub, bạn sẽ thấy file mẫu tên là .env.example. "
        "Hãy sao chép file này thành file .env và cập nhật thông tin:"
    )
    add_code_block(doc, 'cd C:\\Users\\<TenUser>\\hr-ai-project\\WrenAI\\docker\n\n# Tạo file .env\nCopy-Item .env.example -Destination .env\n# Sau đó mở file .env bằng trình soạn thảo và điền GEMINI_API_KEY')

    doc.add_heading("6.2. Nội dung quan trọng trong .env", level=2)
    doc.add_paragraph("File .env chứa các biến cấu hình sau (đã được thiết lập sẵn):")
    add_table(doc,
        ["Biến", "Giá trị", "Mô tả"],
        [
            ["GEMINI_API_KEY", "(nhập key của bạn)", "API key cá nhân của bạn"],
            ["GENERATION_MODEL", "gemini/gemini-2.5-flash", "Model LLM sử dụng"],
            ["COMPOSE_PROJECT_NAME", "wrenai", "Tên Docker project"],
            ["SHOULD_FORCE_DEPLOY", "1", "Tự động deploy model khi khởi động"],
            ["USER_UUID", "(đã có sẵn)", "UUID người dùng — giữ nguyên"],
        ],
        col_widths=[4, 5, 7]
    )
    add_warning_box(doc, "Tuyệt đối không để lộ GEMINI_API_KEY bằng cách commit file .env lên public repo.")

    doc.add_heading("6.3. Cấu hình AI models (config.yaml)", level=2)
    doc.add_paragraph("File config.yaml cũng đã có sẵn, cấu hình 2 model AI:")
    add_table(doc,
        ["Thành phần", "Model", "Vai trò"],
        [
            ["LLM", "gemini/gemini-2.5-flash", "Hiểu câu hỏi tiếng Việt, tạo SQL, phân tích kết quả"],
            ["Embedder", "gemini/gemini-embedding-001", "Tạo vector embeddings cho semantic search (768 dims)"],
        ],
        col_widths=[3, 5, 8]
    )
    add_important_box(doc, "Không cần chỉnh sửa config.yaml. Mọi cấu hình đã sẵn sàng sau khi pull code.")

    doc.add_heading("6.4. Kiểm tra cấu hình", level=2)
    add_code_block(doc, '# Kiểm tra file .env có tồn tại và đầy đủ\nSelect-String -Path .env -Pattern "GEMINI_API_KEY|GENERATION_MODEL|USER_UUID"\n\n# Kiểm tra config.yaml\nTest-Path config.yaml\n# Kết quả mong đợi: True')
    doc.add_paragraph("Kết quả mong đợi: 3 dòng với giá trị thực tế và config.yaml tồn tại.")

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 7: KHỞI ĐỘNG HỆ THỐNG
    # =========================================================================
    doc.add_heading("7. Khởi động hệ thống", level=1)

    doc.add_heading("7.1. Đảm bảo Docker Desktop đang chạy", level=2)
    doc.add_paragraph(
        "Kiểm tra icon Docker ở góc phải taskbar (system tray) phải có màu xanh. "
        "Nếu icon màu cam hoặc không thấy, mở Docker Desktop và đợi khởi động xong."
    )

    doc.add_heading("7.2. Khởi động tất cả dịch vụ", level=2)
    add_step(doc, 1, "Di chuyển vào thư mục docker")
    add_code_block(doc, "cd C:\\Users\\<TenUser>\\hr-ai-project\\WrenAI\\docker")
    add_step(doc, 2, "Khởi động hệ thống (lần đầu sẽ tải images, mất 5–10 phút)")
    add_code_block(doc, "docker compose up -d")

    doc.add_paragraph("Kết quả mong đợi:")
    add_code_block(doc,
        '[+] Running 6/6\n'
        ' ✔ Container wrenai-qdrant-1            Started\n'
        ' ✔ Container wrenai-wren-engine-1       Started\n'
        ' ✔ Container wrenai-ibis-server-1       Started\n'
        ' ✔ Container wrenai-wren-ai-service-1   Started\n'
        ' ✔ Container wrenai-bootstrap-1         Started\n'
        ' ✔ Container wrenai-wren-ui-1           Started'
    )

    doc.add_heading("7.3. Kiểm tra tất cả container đang chạy", level=2)
    add_code_block(doc, 'docker ps --format "table {{.Names}}\\t{{.Status}}\\t{{.Ports}}"')

    add_table(doc,
        ["Container", "Trạng thái mong đợi", "Port"],
        [
            ["wrenai-wren-ui-1", "Up", "0.0.0.0:3000 → 3000"],
            ["wrenai-wren-ai-service-1", "Up", "0.0.0.0:5555 → 5555"],
            ["wrenai-wren-engine-1", "Up", "0.0.0.0:8080 → 8080"],
            ["wrenai-qdrant-1", "Up", "0.0.0.0:6333 → 6333"],
            ["wrenai-ibis-server-1", "Up", "0.0.0.0:8000 → 8000"],
            ["wrenai-bootstrap-1", "Exited (0)", "—"],
        ],
        col_widths=[5, 3.5, 5]
    )

    doc.add_paragraph(
        "Lưu ý: container bootstrap sẽ có trạng thái Exited (0) sau khi khởi tạo xong. "
        "Đây là hành vi bình thường."
    )

    doc.add_heading("7.4. Đợi AI Service sẵn sàng", level=2)
    doc.add_paragraph("AI Service cần 1–2 phút để khởi động hoàn tất. Kiểm tra bằng lệnh:")
    add_code_block(doc, 'Invoke-RestMethod -Uri "http://localhost:5555/health" -Method Get')
    doc.add_paragraph('Kết quả mong đợi: trường status hiển thị giá trị "ok".')
    add_important_box(doc, 'Nếu nhận được lỗi "Unable to connect", đợi thêm 30 giây rồi thử lại.')

    doc.add_heading("7.5. Dừng hệ thống", level=2)
    add_code_block(doc,
        '# Dừng tất cả dịch vụ\n'
        'docker compose down\n'
        '\n'
        '# Dừng và xóa toàn bộ dữ liệu (chỉ khi muốn reset hoàn toàn):\n'
        '# docker compose down -v'
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 8: KIỂM TRA HỆ THỐNG
    # =========================================================================
    doc.add_heading("8. Kiểm tra hệ thống hoạt động", level=1)

    doc.add_heading("8.1. Mở giao diện web", level=2)
    doc.add_paragraph("Mở trình duyệt và truy cập:")
    add_code_block(doc, "http://localhost:3000")

    doc.add_heading("8.2. Onboarding lần đầu", level=2)
    doc.add_paragraph(
        "Nếu là lần đầu chạy trên máy của bạn, hệ thống sẽ hiển thị trang Onboarding — "
        "kết nối database. Thực hiện như sau:"
    )
    add_step(doc, 1, "Chọn loại database: SQL Server")
    add_step(doc, 2, "Nhập thông tin kết nối SQL Server (hỏi quản trị viên dự án)")
    add_step(doc, 3, "Chọn các bảng cần sử dụng theo hướng dẫn")
    add_step(doc, 4, "Cấu hình relationships giữa các bảng")
    add_step(doc, 5, "Hệ thống sẽ tự động deploy model")

    doc.add_paragraph(
        "Nếu hệ thống đã được cấu hình trước (dữ liệu được lưu trong thư mục data/), "
        "giao diện sẽ hiển thị ngay trang Home với ô nhập câu hỏi."
    )

    doc.add_heading("8.3. Thử nghiệm câu hỏi mẫu", level=2)
    doc.add_paragraph('Nhập một trong các câu hỏi sau vào ô "Ask a question":')
    add_table(doc,
        ["Câu hỏi mẫu", "Kết quả mong đợi"],
        [
            ["10 nhân viên có nguy cơ nghỉ việc cao nhất", "Bảng 10 nhân viên với tên, phòng ban, xác suất, risk level"],
            ["Tỷ lệ nghỉ việc theo phòng ban", "Bảng thống kê tỷ lệ nghỉ việc từng phòng ban"],
            ["So sánh lương trung bình giữa nam và nữ", "Bảng so sánh với cột Gender, avg salary"],
            ["Bao nhiêu nhân viên làm thêm giờ", "Số lượng và tỷ lệ nhân viên overtime"],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading("8.4. Kiểm tra biểu đồ (chart)", level=2)
    doc.add_paragraph(
        "Sau khi có kết quả SQL, nhấn nút Chart ở góc phải để xem biểu đồ trực quan. "
        "Hệ thống hỗ trợ: biểu đồ cột (bar chart), biểu đồ đường (line chart), "
        "biểu đồ tròn (donut chart), biểu đồ cột xếp chồng (stacked bar)."
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 9: HƯỚNG DẪN SỬ DỤNG GIAO DIỆN
    # =========================================================================
    doc.add_heading("9. Hướng dẫn sử dụng giao diện Wren AI", level=1)

    doc.add_heading("9.1. Trang Home — đặt câu hỏi", level=2)
    doc.add_paragraph(
        "Đây là trang chính, nơi bạn nhập câu hỏi tiếng Việt. Hệ thống sẽ phân tích câu hỏi, "
        "tìm kiếm SQL Pairs liên quan, tạo câu lệnh SQL, thực thi và trả về kết quả."
    )

    doc.add_heading("9.2. Trang Modeling — xem mô hình dữ liệu", level=2)
    doc.add_paragraph(
        "Hiển thị danh sách các bảng (models) và view, quan hệ giữa các bảng (relationships), "
        "mô tả từng cột (metadata). Đây là nơi quản lý semantic layer."
    )

    doc.add_heading("9.3. Trang Knowledge — SQL Pairs và Instructions", level=2)
    doc.add_paragraph(
        "SQL Pairs: hệ thống đã có 18 cặp câu hỏi — SQL mẫu làm ví dụ cho AI. "
        "Instructions: 13 quy tắc nghiệp vụ (đặt tên, xử lý NULL, format số, điều kiện lọc...) "
        "giúp AI tạo SQL chuẩn xác hơn."
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 10: QUY TRÌNH LÀM VIỆC VỚI GIT
    # =========================================================================
    doc.add_heading("10. Quy trình làm việc với Git", level=1)

    doc.add_heading("10.1. Nguyên tắc vàng — quy tắc bắt buộc", level=2)
    add_error_box(doc, "Không bao giờ push trực tiếp lên branch main.")
    add_important_box(doc, "Luôn làm việc trên branch hr_domain_research.")
    add_important_box(doc, "Tạo Merge Request khi muốn đưa code vào main.")

    doc.add_paragraph(
        'Branch main đã được bảo vệ (protected). Hệ thống sẽ tự chối mọi lệnh push '
        'trực tiếp lên main từ tài khoản Developer. Chỉ Maintainer (quản trị viên) mới '
        'có quyền merge vào main.'
    )

    doc.add_heading("10.2. Quy trình làm việc hàng ngày", level=2)
    doc.add_paragraph("Quy trình 5 bước cho mỗi ngày làm việc:")

    add_table(doc,
        ["Bước", "Hành động", "Lệnh"],
        [
            ["1", "Cập nhật code mới nhất", "git pull origin hr_domain_research"],
            ["2", "Làm việc và test", "(chỉnh sửa code, chạy docker compose up -d)"],
            ["3", "Xem file thay đổi", "git status"],
            ["4", "Commit thay đổi", "git add . && git commit -m \"mô tả\""],
            ["5", "Push lên GitLab", "git push origin hr_domain_research"],
        ],
        col_widths=[1.5, 5, 9]
    )

    doc.add_heading("10.3. Quy tắc viết commit message", level=2)
    add_table(doc,
        ["Tiền tố", "Ý nghĩa", "Ví dụ"],
        [
            ["feat:", "Thêm tính năng mới", 'feat: thêm 5 SQL Pairs cho phòng ban'],
            ["fix:", "Sửa lỗi", 'fix: sửa lỗi query NULL values'],
            ["docs:", "Cập nhật tài liệu", 'docs: cập nhật hướng dẫn sử dụng'],
            ["refactor:", "Tái cấu trúc code", 'refactor: tối ưu config embedder'],
        ],
        col_widths=[3, 4, 8]
    )

    doc.add_heading("10.4. Tạo Merge Request", level=2)
    doc.add_paragraph("Khi code đã sẵn sàng để đưa vào main:")
    add_step(doc, 1, "Truy cập trang tạo Merge Request")
    add_code_block(doc, f"{GITLAB_URL}/-/merge_requests/new")
    add_step(doc, 2, "Chọn source branch: hr_domain_research, target branch: main")
    add_step(doc, 3, "Điền title mô tả thay đổi, description chi tiết")
    add_step(doc, 4, "Assign reviewer là quản trị viên dự án")
    add_step(doc, 5, "Nhấn Create merge request")
    doc.add_paragraph("Quản trị viên sẽ review và merge code của bạn vào main.")

    doc.add_heading("10.5. Xử lý xung đột (merge conflicts)", level=2)
    doc.add_paragraph("Nếu khi pull mà gặp xung đột:")
    add_code_block(doc,
        '# Xem file bị xung đột\n'
        'git status\n'
        '\n'
        '# Mở file bị xung đột, tìm các dòng:\n'
        '# <<<<<<< HEAD\n'
        '# (code của bạn)\n'
        '# =======\n'
        '# (code từ remote)\n'
        '# >>>>>>> origin/hr_domain_research\n'
        '\n'
        '# Sửa file: giữ lại phần code đúng, xóa các dòng <<<, ===, >>>\n'
        '\n'
        '# Sau khi sửa xong\n'
        'git add .\n'
        'git commit -m "fix: resolve merge conflict"\n'
        'git push origin hr_domain_research'
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 11: CẤU TRÚC THƯ MỤC
    # =========================================================================
    doc.add_heading("11. Cấu trúc thư mục dự án", level=1)

    add_code_block(doc,
        'hr-ai-project/\n'
        '│\n'
        '├── .gitignore                          ← Danh sách file không commit lên Git\n'
        '├── ONBOARDING_GUIDE.md                 ← Tài liệu onboarding (Markdown)\n'
        '├── TAI_LIEU_DU_AN_HR_ANALYTICS.md      ← Tài liệu kỹ thuật chi tiết\n'
        '├── MAINTAINER_GUIDE.md                 ← Hướng dẫn quản trị (chỉ dành cho lead)\n'
        '│\n'
        '├── WrenAI/                             ← Engine chính của hệ thống\n'
        '│   └── docker/\n'
        '│       ├── .env                        ← Cấu hình + API key (đã commit)\n'
        '│       ├── .env.example                ← File mẫu tham khảo\n'
        '│       ├── config.yaml                 ← Cấu hình LLM và Embedder\n'
        '│       ├── docker-compose.yaml         ← Cấu hình Docker containers\n'
        '│       └── data/                       ← Dữ liệu runtime (không commit)\n'
        '│\n'
        '├── legacy/                             ← SQL scripts cũ (tham khảo)\n'
        '│   ├── init-db.sql                     ← Tạo bảng và nhập dữ liệu\n'
        '│   ├── create_actionable_views.sql     ← Tạo view phân tích\n'
        '│   └── setup_db_mail_template.sql      ← Cấu hình email cảnh báo\n'
        '│\n'
        '└── notebooks/                          ← Jupyter notebooks\n'
        '    ├── HR_Analytics_Project_Final.ipynb ← Mô hình dự báo nghỉ việc\n'
        '    ├── WA_Fn-UseC_-HR-Employee-Attrition.csv\n'
        '    └── README.md                       ← Hướng dẫn chạy notebook'
    )

    doc.add_heading("11.1. Các file quan trọng cần biết", level=2)
    add_table(doc,
        ["File", "Mức độ", "Mô tả"],
        [
            ["WrenAI/docker/.env", "Cấu hình", "Chứa GEMINI_API_KEY — tuyệt đối không commit file này"],
            ["WrenAI/docker/.env.example", "Tham khảo", "Mẫu file .env để copy ra file thực tế"],
            ["WrenAI/docker/docker-compose.yaml", "Cấu hình", "Định nghĩa 6 Docker services"],
            ["WrenAI/docker/config.yaml", "Cấu hình", "Model AI (Gemini 2.5 Flash), embedder (Gemini Embedding 001), pipelines"],
            ["TAI_LIEU_DU_AN_HR_ANALYTICS.md", "Tài liệu", "Mô tả kỹ thuật chi tiết dự án"],
        ],
        col_widths=[5.5, 2.5, 7]
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 12: XỬ LÝ LỖI THƯỜNG GẶP
    # =========================================================================
    doc.add_heading("12. Xử lý lỗi thường gặp", level=1)

    errors = [
        {
            "title": '12.1. Lỗi: "Cannot connect to Docker daemon"',
            "cause": "Docker Desktop chưa chạy.",
            "fix": '# Mở Docker Desktop\nStart-Process "C:\\Program Files\\Docker\\Docker\\Docker Desktop.exe"\n\n# Đợi 30–60 giây cho Docker khởi động\n# Kiểm tra lại\ndocker ps',
        },
        {
            "title": '12.2. Lỗi: "Port already in use"',
            "cause": "Có ứng dụng khác đang dùng cổng 3000, 5555, 8080...",
            "fix": '# Tìm ứng dụng đang dùng cổng 3000\nnetstat -ano | findstr :3000\n\n# Kết thúc process đó (thay PID bằng số thực tế)\ntaskkill /PID <PID> /F\n\n# Khởi động lại\ncd WrenAI\\docker\ndocker compose up -d',
        },
        {
            "title": '12.3. Lỗi: AI Service trả về lỗi xác thực API',
            "cause": "GEMINI_API_KEY trong file .env bị thiếu, sai, hoặc đã hết quota.",
            "fix": '# Kiểm tra API key trong .env\nSelect-String -Path .env -Pattern "GEMINI_API_KEY"\n# Đảm bảo có giá trị: GEMINI_API_KEY=AIzaSy...\n\n# Khởi động lại:\ncd WrenAI\\docker\ndocker compose down\ndocker compose up -d\n\n# Nếu vẫn lỗi, liên hệ Lead để kiểm tra API key',
        },
        {
            "title": '12.4. Lỗi: "git push rejected" khi push lên main',
            "cause": "Bạn đang cố push lên branch main (bị bảo vệ).",
            "fix": '# Chuyển sang branch đúng\ngit checkout hr_domain_research\n\n# Push lên branch đúng\ngit push origin hr_domain_research',
        },
        {
            "title": "12.5. Lỗi: container wren-ai-service restart liên tục",
            "cause": "GEMINI_API_KEY thiếu/sai, hoặc config.yaml bị lỗi cấu trúc.",
            "fix": '# Xem logs chi tiết\ndocker logs wrenai-wren-ai-service-1 --tail 50\n\n# Kiểm tra .env có GEMINI_API_KEY\nSelect-String -Path WrenAI\\docker\\.env -Pattern "GEMINI_API_KEY"\n\n# Kiểm tra config.yaml tồn tại\nTest-Path WrenAI\\docker\\config.yaml\n\n# Khởi động lại\ncd WrenAI\\docker\ndocker compose down\ndocker compose up -d',
        },
        {
            "title": '12.6. Lỗi: "Authentication failed" khi push code',
            "cause": "Personal Access Token sai hoặc hết hạn.",
            "fix": '# Xóa credential cũ\ncmdkey /delete:git:https://gitlab.com\n\n# Push lại, nhập token mới khi được hỏi\ngit push origin hr_domain_research',
        },
    ]

    for err in errors:
        doc.add_heading(err["title"], level=2)
        p = doc.add_paragraph()
        run = p.add_run("Nguyên nhân: ")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(err["cause"])
        run2.font.size = Pt(11)

        p2 = doc.add_paragraph()
        run3 = p2.add_run("Cách xử lý:")
        run3.bold = True
        run3.font.size = Pt(11)
        add_code_block(doc, err["fix"])

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 13: QUY TẮC BẢO MẬT
    # =========================================================================
    doc.add_heading("13. Quy tắc bảo mật bắt buộc", level=1)

    doc.add_heading("13.1. Các file nhạy cảm và quy tắc", level=2)
    doc.add_paragraph(
        "File .env chứa Gemini API key đã được commit vào GitLab private repo để team "
        "pull về và chạy ngay. Điều này an toàn vì repo là private — chỉ thành viên được mời "
        "mới truy cập được."
    )
    add_table(doc,
        ["File", "Trạng thái", "Lý do"],
        [
            [".env", "Đã commit (private repo)", "Chứa API key dùng chung — nhúng sẵn cho team"],
            ["config.yaml", "Đã commit", "Cấu hình AI models — nhúng sẵn cho team"],
            ["gcloud/ folder", "Không commit (.gitignore)", "Credentials cá nhân — không dùng nữa"],
            ["service-account*.json", "Không commit", "Khóa tài khoản dịch vụ — không cần"],
            ["File chứa mật khẩu SQL Server", "Không commit", "Bảo vệ truy cập database"],
        ],
        col_widths=[5, 4, 7]
    )

    doc.add_heading("13.2. Kiểm tra trước khi commit", level=2)
    add_code_block(doc,
        '# Luôn kiểm tra trước khi commit\n'
        'git status\n'
        '\n'
        '# Đảm bảo không có file credentials cá nhân hoặc file lạ\n'
        '# Nếu thấy file .json chứa credentials, dừng lại và liên hệ Lead'
    )

    doc.add_heading("13.3. Nếu lỡ commit file bí mật", level=2)
    add_code_block(doc,
        '# Xóa file credentials cá nhân khỏi Git (giữ file trên máy)\n'
        'git rm --cached <tên-file-credentials>\n'
        'git commit -m "fix: xóa file credentials khỏi git"\n'
        'git push origin hr_domain_research\n'
        '\n'
        '# Nếu lỡ commit API key khác lên public repo:\n'
        '# Revoke key ngay trên Google AI Studio và tạo key mới'
    )

    doc.add_heading("13.4. Bảo vệ API key cá nhân", level=2)
    doc.add_paragraph(
        "Bạn chịu trách nhiệm bảo quản API key cá nhân của mình. "
        "Không chia sẻ API key cho ai khác, không post lên forum/chat công khai."
    )
    add_error_box(doc, "KHÔNG chia sẻ GEMINI_API_KEY và tuyệt đối không push file .env lên public repo.")
    add_important_box(doc, "Nếu API key bị lộ: truy cập ngay Google AI Studio để revoke key cũ và tạo key mới.")

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 14: CHECKLIST ONBOARDING
    # =========================================================================
    doc.add_heading("14. Bảng kiểm (checklist) onboarding", level=1)
    doc.add_paragraph("Đánh dấu ✓ khi hoàn thành từng bước:")

    checklist_items = [
        ["1", "Đã cài Docker Desktop và chạy thành công", "☐"],
        ["2", "Đã cài Git và cấu hình user.name, user.email", "☐"],
        ["3", "Đã nhận quyền Developer trên GitLab (chấp nhận lời mời)", "☐"],
        ["4", "Đã tạo Personal Access Token trên GitLab", "☐"],
        ["5", "Đã clone dự án thành công", "☐"],
        ["6", "Đã checkout branch hr_domain_research", "☐"],
        ["7", "Đã kiểm tra file .env tồn tại (Test-Path .env → True)", "☐"],
        ["8", "Đã kiểm tra file config.yaml tồn tại", "☐"],
        ["9", "Đã chạy docker compose up -d thành công", "☐"],
        ["10", "Đã kiểm tra 6 container đang chạy (docker ps)", "☐"],
        ["11", "Đã kiểm tra AI Service health (localhost:5555/health → ok)", "☐"],
        ["12", "Đã mở giao diện web tại localhost:3000", "☐"],
        ["13", "Đã thử đặt 1 câu hỏi và nhận được kết quả", "☐"],
        ["14", "Đã thử tạo 1 commit và push lên hr_domain_research", "☐"],
    ]

    add_table(doc,
        ["STT", "Hạng mục", "Trạng thái"],
        checklist_items,
        col_widths=[1.5, 11, 2.5]
    )

    add_important_box(doc, "Nếu hoàn thành tất cả 14 bước trên, bạn đã sẵn sàng làm việc! Chỉ cần 3 phần mềm: Docker, Git, trình duyệt web.")

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 15: PHÂN CÔNG NHIỆM VỤ CHI TIẾT THEO TEAM
    # =========================================================================
    doc.add_heading("15. Phân công nhiệm vụ chi tiết theo team", level=1)
    doc.add_paragraph(
        "Dự án được chia thành 3 nhóm công việc chính, mỗi nhóm có nhiệm vụ cụ thể "
        "cho buổi Showcase vấn đáp dự án tuần 10/2 – 16/2/2026. "
        "Mỗi thành viên cần hoàn thành đầy đủ phần Research + Deep Dive + Output deliverables."
    )

    # --- Bảng tổng quan phân công ---
    doc.add_heading("15.1. Bảng tổng quan phân công", level=2)
    add_table(doc,
        ["Field", "Người phụ trách", "Công việc chính", "Deadline"],
        [
            ["AI Engineering", "Khải (Lead), Hân, Ninh", "Build Semantic Layer + Deep Dive Wren AI", "16/02/2026"],
            ["Data Analytics", "Gia", "Analyze Business Insights + Notebook Deep Dive", "16/02/2026"],
            ["Data Pipeline & MLOps", "Uyên", "ETL flow + Auto-MLOps Workflow Deep Dive", "16/02/2026"],
        ],
        col_widths=[3.5, 4, 6, 2.5]
    )

    doc.add_page_break()

    # =========================================================================
    # 15.2 AI ENGINEERING — KHẢI, HÂN, NINH
    # =========================================================================
    doc.add_heading("15.2. AI Engineering — Khải (Lead), Hân, Ninh", level=2)

    doc.add_heading("I. CÔNG VIỆC CHÍNH", level=3)

    doc.add_paragraph("Công việc 1: Chạy lại toàn bộ hệ thống")
    p = doc.add_paragraph()
    run = p.add_run(
        "Mỗi thành viên tự cài đặt và chạy toàn bộ hệ thống trên máy cá nhân theo hướng dẫn "
        "Chương 1–8 của tài liệu này. Đảm bảo 6 Docker containers hoạt động, AI Service health "
        "trả về status: ok, giao diện web tại localhost:3000 hoạt động."
    )
    run.font.size = Pt(10)

    doc.add_paragraph("Công việc 2: Mỗi thành viên tạo 5 nghiệp vụ HR mới")
    p = doc.add_paragraph()
    run = p.add_run(
        "Mỗi người tạo 5 business rules mới trên giao diện Wren AI (localhost:3000). "
        "Mỗi nghiệp vụ phải bao gồm đầy đủ 4 thành phần:"
    )
    run.font.size = Pt(10)

    add_table(doc,
        ["Thành phần", "Mô tả", "Nơi cấu hình trên Wren AI"],
        [
            ["Relations", "Thiết lập quan hệ giữa các bảng (JOIN conditions)", "Trang Modeling → chọn bảng → Add Relationship"],
            ["Modelling", "Thêm calculated fields, measures, dimensions vào semantic layer", "Trang Modeling → chọn bảng → thêm cột"],
            ["SQL Pairs", "Cặp câu hỏi tiếng Việt – SQL mẫu để AI học theo", "Trang Knowledge → SQL Pairs → Add"],
            ["Instructions", "Quy tắc nghiệp vụ bằng text (ví dụ: điều kiện lọc, format kết quả)", "Trang Knowledge → Instructions → Add"],
        ],
        col_widths=[2.5, 6, 7]
    )

    doc.add_paragraph("Yêu cầu cho mỗi nghiệp vụ:")
    items = [
        "Phải test trên Wren AI: đặt câu hỏi tiếng Việt → hệ thống trả SQL đúng → kết quả chính xác",
        "Chụp screenshot kết quả chạy thành công làm minh chứng",
        "Ghi lại: tên nghiệp vụ, mô tả, SQL mẫu, câu hỏi test, kết quả",
    ]
    for item in items:
        p = doc.add_paragraph(f"• {item}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_paragraph("Phân chia 5 nghiệp vụ gợi ý theo từng người:")
    add_table(doc,
        ["Thành viên", "5 nghiệp vụ HR gợi ý", "Ghi chú"],
        [
            ["Khải (Lead)", "1. Phân tích overtime theo phòng ban và level\\n2. So sánh lương trung bình theo Education Field\\n3. Top nhân viên có years at company cao nhất theo department\\n4. Tỷ lệ nhân viên có work-life balance thấp theo job role\\n5. Phân tích mối quan hệ giữa training times và performance rating", "Lead review toàn bộ nghiệp vụ của Hân và Ninh"],
            ["Hân", "1. Nhân viên có distance from home > 20km theo department\\n2. Tỷ lệ attrition theo marital status và gender\\n3. Phân tích monthly income theo job level và years in current role\\n4. Nhân viên có stock option level = 0 và risk cao\\n5. So sánh job satisfaction giữa overtime và non-overtime", "Cần verify với Khải trước khi submit"],
            ["Ninh", "1. Top 10 nhân viên có percent salary hike thấp nhất\\n2. Phân tích relationship satisfaction theo age group\\n3. Tỷ lệ attrition theo number of companies worked\\n4. Nhân viên có total working years > 20 và performance rating thấp\\n5. So sánh environment satisfaction giữa các department", "Cần verify với Khải trước khi submit"],
        ],
        col_widths=[2.5, 10, 3.5]
    )

    doc.add_paragraph("Công việc 3: Tài liệu Deep Dive kỹ thuật")
    p = doc.add_paragraph()
    run = p.add_run(
        "Chuẩn bị trả lời các câu hỏi Deep Dive về source code và kiến trúc Wren AI. "
        "Phân chia nghiên cứu theo từng người:"
    )
    run.font.size = Pt(10)

    doc.add_heading("Phân chia Deep Dive — Khải (Lead):", level=3)
    deep_dive_khai = [
        "Kiến trúc tổng thể Wren AI: 6 services, luồng dữ liệu từ câu hỏi → SQL → kết quả",
        "Tại sao chọn Wren AI thay vì LangChain làm nòng cốt? (so sánh: semantic layer vs prompt engineering)",
        "Semantic Layer giải quyết bài toán gì cho Text-to-SQL? Tại sao cần metadata thay vì gửi trực tiếp schema?",
        "Folder source code chính: wren-ai-service/src/pipelines/ — giải thích từng pipeline: sql_generation, sql_correction, intent_classification",
        "Review và tổng hợp toàn bộ Deep Dive output của Hân và Ninh",
    ]
    for i, item in enumerate(deep_dive_khai, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_heading("Phân chia Deep Dive — Hân:", level=3)
    deep_dive_han = [
        "Cơ chế bảo mật: làm sao bảo mật thông tin nhân viên khi query qua LLM? (Semantic Layer che giấu schema thật)",
        "Docker Compose architecture: 6 services giao tiếp thế nào? Network, ports, volumes",
        "Config.yaml: ý nghĩa từng section (llm, embedder, document_store, pipeline, settings)",
        "Qdrant vector database: vai trò lưu trữ embeddings cho SQL Pairs và Instructions",
        "File wren-ai-service/src/providers/ — phân tích LLM provider (LiteLLM), Embedder provider, Document Store provider",
    ]
    for i, item in enumerate(deep_dive_han, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_heading("Phân chia Deep Dive — Ninh:", level=3)
    deep_dive_ninh = [
        "Luồng kết nối SQL Server: Wren Engine → Ibis Server → SQL Server. Giải thích API endpoint và connection flow",
        "Wren UI (Next.js frontend): cấu trúc src/, API routes, GraphQL schema",
        "Wren MDL (Model Definition Language): schema JSON, cách định nghĩa models/relationships/metrics",
        "Context cần cung cấp cho AI để hiểu nghiệp vụ HR: Instructions, SQL Pairs, table descriptions — cơ chế RAG",
        "File wren-ai-service/src/pipelines/retrieval/ — cơ chế retrieval: db_schema_retrieval, sql_pairs_retrieval, instructions_retrieval",
    ]
    for i, item in enumerate(deep_dive_ninh, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_heading("II. THROUGH BACK — Câu hỏi vấn đáp cần chuẩn bị", level=3)
    doc.add_paragraph("Toàn bộ team AI Engineering cần chuẩn bị trả lời các câu hỏi sau:")
    throughback_questions = [
        "Tại sao chọn Wren AI thay vì LangChain nòng cốt? So sánh ưu/nhược điểm?",
        "Làm sao để bảo mật thông tin nhân viên khi query qua LLM?",
        "Semantic Layer trong Wren AI hoạt động thế nào? Giải quyết bài toán gì?",
        "Các tính năng chính là gì, sắp xếp trong folder nào, chạy class nào để call?",
        "Làm sao kết nối tới SQL Server qua gì (API)? Viết truy vấn và đảm bảo đúng để thực thi SQL?",
        "Context nào cần cung cấp cho AI để nó hiểu quy trình nghiệp vụ HR mới thêm vào?",
        "Repo chạy với file cấu hình metadata đúng chuẩn như thế nào?",
        "Demo live: đặt câu hỏi tiếng Việt → AI trả SQL → kết quả đúng với nghiệp vụ mới tạo",
    ]
    for i, q in enumerate(throughback_questions, 1):
        p = doc.add_paragraph(f"{i}. {q}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_heading("III. SOURCE CODE CẦN NGHIÊN CỨU", level=3)
    add_table(doc,
        ["File/Folder", "Người phụ trách", "Nội dung cần nắm"],
        [
            ["WrenAI/docker/docker-compose.yaml", "Hân", "Kiến trúc 6 services, ports, volumes, env vars"],
            ["WrenAI/docker/config.yaml", "Hân", "Cấu hình LLM (gemini/gemini-2.5-flash), embedder (gemini-embedding-001), pipelines"],
            ["WrenAI/wren-ai-service/src/pipelines/", "Khải", "Pipeline SQL generation, correction, retrieval"],
            ["WrenAI/wren-ai-service/src/providers/", "Hân", "LLM/Embedder/DocStore providers implementation"],
            ["WrenAI/wren-ai-service/src/globals.py", "Khải", "Service container initialization, pipe components"],
            ["WrenAI/wren-ai-service/src/__main__.py", "Khải", "Application entrypoint, lifespan, FastAPI setup"],
            ["WrenAI/wren-ui/src/", "Ninh", "Next.js frontend: pages, components, API routes"],
            ["WrenAI/wren-mdl/mdl.schema.json", "Ninh", "MDL schema: models, relationships, metrics definition"],
            ["WrenAI/wren-engine/", "Ninh", "Query engine: SQL execution, connection management"],
            ["legacy/init-db.sql", "Cả team", "Database schema gốc: bảng, cột, dữ liệu mẫu"],
            ["legacy/create_actionable_views.sql", "Cả team", "Views phân tích: risk level, burnout, retention metrics"],
        ],
        col_widths=[5.5, 2.5, 8]
    )

    doc.add_heading("IV. OUTPUT DELIVERABLES", level=3)
    add_table(doc,
        ["Output", "Người phụ trách", "Chi tiết"],
        [
            ["15 nghiệp vụ HR mới (5/người)", "Khải, Hân, Ninh", "Mỗi nghiệp vụ: relation + modelling + SQL pair + instruction"],
            ["Demo live Wren AI", "Cả team", "Chạy 5 câu hỏi nghiệp vụ mới → AI trả SQL đúng + kết quả"],
            ["Tài liệu Deep Dive kỹ thuật", "Khải tổng hợp", "Trả lời 8 câu hỏi Through Back + phân tích source code"],
            ["Screenshot minh chứng", "Mỗi người", "Chụp kết quả chạy nghiệp vụ mới trên Wren AI"],
        ],
        col_widths=[4.5, 3, 8.5]
    )

    doc.add_page_break()

    # =========================================================================
    # 15.3 DATA ANALYTICS — GIA
    # =========================================================================
    doc.add_heading("15.3. Data Analytics — Gia", level=2)

    doc.add_heading("I. CÔNG VIỆC CHÍNH", level=3)

    doc.add_paragraph("Công việc 1: Trực quan hóa và giải thích Model")
    items_gia_1 = [
        "Trực quan hóa kết quả 'tr_attrition_result': feature importance chart, confusion matrix",
        "Giải thích mô hình Random Forest: cách hoạt động, hyperparameters, cross-validation",
        "Phân tích: Department nào có Risk cao nhất? Tại sao? (dùng dữ liệu thực từ dataset)",
        "Giải thích Top 3 Feature drivers thực tế từ dữ liệu (OverTime, MonthlyIncome, Age...)",
    ]
    for i, item in enumerate(items_gia_1, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph("Công việc 2: Giải thích từng cell trong Notebook")
    items_gia_2 = [
        "Đọc và giải thích chi tiết từng cell trong file notebooks/HR_Analytics_Project_Final.ipynb",
        "Bổ sung markdown giải thích cho các cell code chưa có annotation",
        "Đảm bảo notebook chạy end-to-end không lỗi",
    ]
    for i, item in enumerate(items_gia_2, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_heading("II. CÂU HỎI DEEP DIVE CẦN CHUẨN BỊ", level=3)
    deep_dive_gia = [
        "Insight từ model giúp gì cho HR Director ra quyết định?",
        "Chi phí thay thế 1 nhân sự là bao nhiêu? (trích dẫn báo cáo số liệu đáng tin cậy)",
        "Data Leakage là gì? Tại sao tách Train/Test bình thường lại sai trong bài toán này?",
        "OOF (Out-of-Fold) giúp mô phỏng Production như thế nào?",
        "Chỉ số Recall quan trọng hơn Precision không? Tại sao? (trong bài toán dự báo nghỉ việc)",
        "Cách đọc confusion matrix: True Positive, False Negative nghĩa là gì trong ngữ cảnh HR?",
        "Feature importance: OverTime tại sao là yếu tố số 1? Dữ liệu chứng minh như thế nào?",
    ]
    for i, q in enumerate(deep_dive_gia, 1):
        p = doc.add_paragraph(f"{i}. {q}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_heading("III. SOURCE CODE CẦN NGHIÊN CỨU", level=3)
    add_table(doc,
        ["File", "Nội dung cần nắm"],
        [
            ["notebooks/HR_Analytics_Project_Final.ipynb", "Toàn bộ EDA, feature engineering, model training, evaluation"],
            ["notebooks/WA_Fn-UseC_-HR-Employee-Attrition.csv", "Dataset gốc IBM HR: 1470 records, 35 features"],
            ["legacy/init-db.sql", "Schema database, view 'tr_attrition_result' và 'vw_attrition_analysis'"],
            ["legacy/create_actionable_views.sql", "Views phân tích: risk scoring, burnout detection"],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading("IV. OUTPUT DELIVERABLES", level=3)
    add_table(doc,
        ["Output", "Chi tiết"],
        [
            ["DOCX Report Deep Dive", "Trả lời 7 câu hỏi Deep Dive + biểu đồ minh họa + số liệu trích dẫn"],
            ["Notebook hoàn chỉnh", "Giải thích từng cell, chạy end-to-end, output rõ ràng"],
            ["Slide/tài liệu trình bày", "Feature importance chart, confusion matrix, department risk analysis"],
        ],
        col_widths=[5, 11]
    )

    doc.add_page_break()

    # =========================================================================
    # 15.4 DATA PIPELINE & MLOPS — UYÊN
    # =========================================================================
    doc.add_heading("15.4. Data Pipeline & MLOps — Uyên", level=2)

    doc.add_heading("I. CÔNG VIỆC CHÍNH", level=3)

    doc.add_paragraph("Công việc 1: Làm rõ luồng ETL Data và MLOps")
    items_uyen_1 = [
        "Vẽ sơ đồ luồng ETL: Data Source → Validation → Processing → Model → Output → Notification",
        "Giải thích Quy Trình Vận Hành Tự Động (Auto-MLOps Workflow) trong Production",
        "Mô tả chi tiết 5 bước của Scheduled Workflow: Trigger → Validation → Retraining → Logging → Notification",
        "So sánh: Notebook thí nghiệm (PoC) vs Hệ thống AI vận hành (Production)",
    ]
    for i, item in enumerate(items_uyen_1, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph("Công việc 2: Tìm hiểu Production Tools")
    items_uyen_2 = [
        "Apache Airflow: cách quản lý luồng dữ liệu phức tạp (DAGs)",
        "SQL Server Agent Job: cách đồng bộ hóa với Database, scheduled jobs",
        "Local LLM (gpt-oss-max): tại sao dùng local LLM cho bảo mật dữ liệu PII",
    ]
    for i, item in enumerate(items_uyen_2, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_heading("II. CÂU HỎI DEEP DIVE CẦN CHUẨN BỊ", level=3)
    deep_dive_uyen = [
        "Trigger logic: Tại sao chạy theo tháng? Độ trễ dữ liệu là gì?",
        "Data Validation: Cách hệ thống chống lại 'Dữ liệu rác' (Garbage in, Garbage out)?",
        "Model Drift: Tại sao mô hình có thể yếu đi theo thời gian? Khi nào cần Retrain?",
        "Monitoring: Giám sát phân phối dữ liệu để phát hiện bất thường như thế nào?",
        "Human-in-the-loop: Vai trò thực sự của AI Agent trong việc hỗ trợ HR Director ra quyết định?",
        "Tại sao dùng Local LLM? (Bảo mật PII, ISO/GDPR, dữ liệu không rời Server nội bộ)",
        "Giá trị của Email Insight & HTML Report: AI đóng vai trò Analyst chuyên nghiệp như thế nào?",
        "So sánh Apache Airflow vs SQL Server Agent Job vs Prefect cho use case này?",
    ]
    for i, q in enumerate(deep_dive_uyen, 1):
        p = doc.add_paragraph(f"{i}. {q}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    doc.add_heading("III. SOURCE CODE CẦN NGHIÊN CỨU", level=3)
    add_table(doc,
        ["File", "Nội dung cần nắm"],
        [
            ["notebooks/HR_Analytics_Project_Final.ipynb", "Luồng data pipeline trong notebook: load → EDA → train → predict → export"],
            ["legacy/setup_db_mail_template.sql", "Cấu hình email alert template cho DB Mail"],
            ["legacy/init-db.sql", "Schema database, stored procedures, data import flow"],
            ["legacy/create_actionable_views.sql", "Views tạo Actionable Insights cho HR report"],
            ["TAI_LIEU_DU_AN_HR_ANALYTICS.md", "Tài liệu kỹ thuật tổng quan dự án — phần MLOps workflow"],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading("IV. OUTPUT DELIVERABLES", level=3)
    add_table(doc,
        ["Output", "Chi tiết"],
        [
            ["Sơ đồ luồng ETL", "Flowchart/diagram: Data Source → Processing → Model → Output → Alert"],
            ["DOCX Report Deep Dive", "Trả lời 8 câu hỏi Deep Dive + so sánh tools + sơ đồ minh họa"],
            ["Slide/tài liệu trình bày", "Tổng quan Data Pipeline, MLOps workflow, Production deployment"],
        ],
        col_widths=[5, 11]
    )

    doc.add_page_break()

    # =========================================================================
    # 15.5 TIMELINE VÀ QUY TRÌNH BÁO CÁO
    # =========================================================================
    doc.add_heading("15.5. Timeline và quy trình báo cáo", level=2)

    add_table(doc,
        ["Ngày", "Milestone", "Ai làm gì"],
        [
            ["10-11/02", "Setup & Onboarding", "Tất cả: clone repo, cài đặt môi trường, chạy hệ thống"],
            ["11-12/02", "Nghiên cứu source code", "Mỗi người đọc file được phân công trong mục III"],
            ["12-13/02", "Tạo nghiệp vụ HR (AI Eng)", "Khải/Hân/Ninh: mỗi người tạo 5 nghiệp vụ trên Wren AI"],
            ["13-14/02", "Deep Dive research", "Tất cả: nghiên cứu câu hỏi Deep Dive, viết tài liệu"],
            ["14-15/02", "Viết report & chuẩn bị demo", "Gia/Uyên: DOCX report. AI Eng: chuẩn bị demo live"],
            ["15-16/02", "Review & rehearsal", "Khải review toàn bộ, team chạy thử showcase"],
        ],
        col_widths=[2.5, 4, 9.5]
    )

    doc.add_heading("Quy trình báo cáo tiến độ:", level=3)
    report_items = [
        "Mỗi ngày 21:00: mỗi người push commit lên branch hr_domain_research với những gì đã làm",
        "Khải (Lead) tổng hợp tiến độ hàng ngày và báo cáo cho Maintainer",
        "Nếu gặp blocker: tạo Issue trên GitLab ngay lập tức, tag @khainn23406",
        "Trước buổi Showcase: Khải review toàn bộ output, chạy thử demo trước 1 ngày",
    ]
    for item in report_items:
        p = doc.add_paragraph(f"• {item}")
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(10)

    add_warning_box(doc, "Tất cả output phải được commit và push lên branch hr_domain_research trước deadline 16/02/2026.")

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 16: LIÊN HỆ HỖ TRỢ
    # =========================================================================
    doc.add_heading("16. Liên hệ hỗ trợ", level=1)

    doc.add_heading("16.1. Khi gặp vấn đề", level=2)
    doc.add_paragraph(
        "1. Đọc lại tài liệu này — chương 12 (xử lý lỗi thường gặp).\n"
        "2. Đọc tài liệu kỹ thuật — file TAI_LIEU_DU_AN_HR_ANALYTICS.md.\n"
        "3. Tạo Issue trên GitLab: truy cập link bên dưới, mô tả chi tiết lỗi gặp phải, "
        "đính kèm screenshot nếu có.\n"
        "4. Liên hệ quản trị viên dự án (Lead/Maintainer) qua email hoặc chat nhóm."
    )
    add_code_block(doc, f"{GITLAB_URL}/-/issues/new")

    doc.add_heading("16.2. Thông tin dự án", level=2)
    add_table(doc,
        ["Hạng mục", "Chi tiết"],
        [
            ["GitLab Repository", GITLAB_URL],
            ["Branch làm việc", "hr_domain_research"],
            ["Branch chính (protected)", "main"],
            ["Giao diện web (local)", "http://localhost:3000"],
            ["AI Service health check", "http://localhost:5555/health"],
            ["Mã số đề tài", PROJECT_CODE],
            ["LLM Model", "gemini/gemini-2.5-flash"],
            ["Embedder Model", "gemini/gemini-embedding-001"],
        ],
        col_widths=[5, 11]
    )

    # =========================================================================
    # FOOTER
    # =========================================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Hết —")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Tài liệu được tạo bởi nhóm dự án HR Analytics. Mọi quyền được bảo lưu.\n"
                     f"Phiên bản {VERSION} — Ngày {DATE}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    return doc


# ==============================================================================
# MAIN
# ==============================================================================
if __name__ == "__main__":
    print("Đang tạo tài liệu onboarding DOCX...")
    document = create_document()

    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), OUTPUT_FILE)
    document.save(output_path)
    print(f"✅ Đã tạo thành công: {output_path}")
    print(f"   Kích thước: {os.path.getsize(output_path) / 1024:.1f} KB")
