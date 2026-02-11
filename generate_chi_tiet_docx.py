#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script sinh tài liệu DOCX chi tiết cho Team HR Analytics AI
Bao gồm: GitLab Access Token, SQL Server Restore, 5 Use Cases, Git Workflow
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_heading_with_color(doc, text, level, color=(0, 102, 204)):
    """Thêm heading với màu tùy chỉnh"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
    return heading

def add_table_border(table):
    """Thêm border cho table"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def set_cell_background(cell, color_code):
    """Đặt màu nền cho cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_code)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_detailed_guide_docx():
    """Tạo tài liệu DOCX chi tiết với access token, SQL setup, 5 use cases, git workflow"""
    doc = Document()
    
    # ===== COVER =====
    title = doc.add_heading('HƯỚNG DẪN ĐẦY ĐỦ TEAM PHÁT TRIỂN', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('HR Analytics AI — Clone GitLab + Setup SQL + 5 Use Cases + Push')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    
    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ('Mã học phần', '252BIM500601'),
        ('GitLab URL', 'https://gitlab.com/boygia757-netizen/hr-ai-project'),
        ('Branch chính', 'hr_domain_research (cấm push main)'),
        ('Public Demo', 'https://certification-lows-spy-tension.trycloudflare.com'),
        ('SQL Database', 'HR_Analytics (restore từ HR_Analytics.bak)'),
        ('Timeline', '11/02 - 16/02/2026'),
        ('Ngày phát hành', datetime.datetime.now().strftime('%d/%m/%Y %H:%M'))
    ]
    
    for idx, (label, value) in enumerate(info_data):
        cells = info_table.rows[idx].cells
        cells[0].text = label
        cells[1].text = value
        set_cell_background(cells[0], 'E7E6E6')
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ===== PHẦN 1: GITLAB ACCESS =====
    add_heading_with_color(doc, '✅ PHẦN 1: GITLAB ACCESS TOKEN & CLONE', level=1)
    
    add_heading_with_color(doc, '1.1 Access Token cho Mỗi Thành Viên', level=2, color=(0, 102, 204))
    
    access_table = doc.add_table(rows=6, cols=4)
    access_table.style = 'Light Grid Accent 1'
    add_table_border(access_table)
    
    access_data = [
        ('Tên', 'Mã SV', 'Token Name', 'Role'),
        ('Khải', 'K234060700', 'khai', 'Developer'),
        ('Hân', 'K234060691', 'han', 'Developer'),
        ('Ninh', 'K234060716', 'ninh', 'Developer'),
        ('Uyên', 'K234060737', 'uyen', 'Developer'),
        ('Gia', 'K234060689', 'gia', 'Developer')
    ]
    
    for idx, row_data in enumerate(access_data):
        cells = access_table.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            for cell_idx, value in enumerate(row_data):
                cells[cell_idx].text = value
    
    doc.add_paragraph()
    
    doc.add_paragraph('Scopes cho tất cả Developer tokens:')
    scopes = [
        'api (truy cập full API)',
        'read_api (đọc repository)',
        'create_runner (tạo CI/CD runners)',
        'write_repository (push code lên)',
        'read_repository (đọc code)',
        'write_registry (push artifacts)'
    ]
    for scope in scopes:
        doc.add_paragraph(scope, style='List Bullet')
    
    add_heading_with_color(doc, '1.2 Clone Repository bằng 1 Câu Lệnh', level=2, color=(0, 102, 204))
    
    doc.add_paragraph('Sử dụng token của bạn (hoặc của gia nếu chưa có token riêng):')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('git clone https://oauth2:[YOUR_TOKEN]@gitlab.com/boygia757-netizen/hr-ai-project.git')
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(192, 0, 0)
    
    p = doc.add_paragraph()
    run = p.add_run('cd hr-ai-project')
    run.font.name = 'Courier New'
    
    p = doc.add_paragraph()
    run = p.add_run('git checkout hr_domain_research')
    run.font.name = 'Courier New'
    
    doc.add_paragraph()
    
    doc.add_paragraph('Ví dụ cụ thể cho Khải (dùng token tạm):')
    p = doc.add_paragraph()
    run = p.add_run('git clone https://oauth2:glpat-I1s2qe7-q09FgrR7nuXxtG86MQp10mtsYzZxCw@gitlab.com/boygia757-netizen/hr-ai-project.git')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_page_break()
    
    # ===== PHẦN 2: SQL SERVER RESTORE =====
    add_heading_with_color(doc, '✅ PHẦN 2: RESTORE SQL SERVER TỪ .BAK', level=1)
    
    add_heading_with_color(doc, '2.1 Bước 1: Copy File .bak vào Docker', level=2, color=(0, 102, 204))
    
    steps = [
        'Từ thư mục gốc dự án, mở PowerShell',
        'Chạy lệnh: docker cp HR_Analytics.bak hr-sql-server:/var/opt/mssql/backup/',
        'Verify file đã copy: docker exec hr-sql-server ls -la /var/opt/mssql/backup/'
    ]
    
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    
    add_heading_with_color(doc, '2.2 Bước 2: Restore Database', level=2, color=(0, 102, 204))
    
    doc.add_paragraph('Kết nối vào SQL Server container qua sqlcmd:')
    p = doc.add_paragraph()
    run = p.add_run('docker exec -it hr-sql-server /opt/mssql-tools/bin/sqlcmd -S localhost -U SA -P "YourSAPassword"')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_paragraph()
    doc.add_paragraph('Chạy lệnh restore (paste vào sqlcmd prompt):')
    
    restore_sql = '''RESTORE DATABASE HR_Analytics 
FROM DISK = '/var/opt/mssql/backup/HR_Analytics.bak'
WITH MOVE 'HR_Analytics' TO '/var/opt/mssql/data/HR_Analytics.mdf',
     MOVE 'HR_Analytics_log' TO '/var/opt/mssql/data/HR_Analytics_log.ldf'
GO'''
    
    p = doc.add_paragraph(restore_sql)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_paragraph()
    doc.add_paragraph('Thoát sqlcmd: EXIT')
    
    add_heading_with_color(doc, '2.3 Bước 3: Verify Restore Thành Công', level=2, color=(0, 102, 204))
    
    verify_cmd = 'docker exec -it hr-sql-server /opt/mssql-tools/bin/sqlcmd -S localhost -U SA -P "YourSAPassword" -Q "SELECT COUNT(*) as EmployeeCount FROM HR_Analytics.dbo.MS_EMPLOYEE"'
    p = doc.add_paragraph(verify_cmd)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_paragraph('Kết quả mong đợi: Hiển thị số lượng nhân viên (VD: 1470)')
    
    doc.add_page_break()
    
    # ===== PHẦN 3: DOCKER STARTUP =====
    add_heading_with_color(doc, '✅ PHẦN 3: KHỞI ĐỘNG DOCKER & VERIFY', level=1)
    
    add_heading_with_color(doc, '3.1 Start Docker Compose', level=2, color=(0, 102, 204))
    
    doc.add_paragraph('Từ thư mục WrenAI/docker:')
    p = doc.add_paragraph('docker compose up -d')
    for run in p.runs:
        run.font.name = 'Courier New'
    
    doc.add_paragraph()
    doc.add_paragraph('Kiểm tra 6 containers UP (chờ ~30-60 giây):')
    p = doc.add_paragraph('docker ps --format "table {{.Names}}\\t{{.Status}}\\t{{.Ports}}"')
    for run in p.runs:
        run.font.name = 'Courier New'
    
    doc.add_paragraph()
    
    containers_check = [
        'wrenai-wren-ui-1 (3000/tcp) — Giao diện chính',
        'wrenai-wren-ai-service-1 (5555/tcp) — AI Service + 29 Pipelines',
        'wrenai-wren-engine-1 (8080/tcp) — MDL Compiler',
        'wrenai-ibis-server-1 (8000/tcp) — SQL Translator',
        'wrenai-qdrant-1 (6333/tcp) — Vector Database',
        'hr-sql-server (1433/tcp) — SQL Server (nếu chạy trong Docker)'
    ]
    
    for check in containers_check:
        doc.add_paragraph(check, style='List Bullet')
    
    add_heading_with_color(doc, '3.2 Health Check', level=2, color=(0, 102, 204))
    
    health_checks = [
        ('Invoke-RestMethod -Uri "http://localhost:5555/health"', '→ Kết quả: {"status":"ok"}'),
        ('Invoke-RestMethod -Uri "http://localhost:6333/health"', '→ Qdrant UP'),
        ('Start-Process "http://localhost:3000"', '→ Mở UI (đợi 30s load)')
    ]
    
    for cmd, result in health_checks:
        p = doc.add_paragraph(f'{cmd}')
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        doc.add_paragraph(result, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== PHẦN 4: 5 USE CASES CHI TIẾT =====
    add_heading_with_color(doc, '✅ PHẦN 4: CẤU HÌNH & TEST 5 NGHIỆP VỤ HR', level=1)
    
    # Khải
    add_heading_with_color(doc, '4.1 KHẢI — 5 Nghiệp Vụ Infrastructure', level=2, color=(204, 102, 0))
    
    khải_uc_table = doc.add_table(rows=6, cols=4)
    khải_uc_table.style = 'Light Grid Accent 1'
    add_table_border(khải_uc_table)
    
    khải_uc = [
        ('STT', 'Nghiệp vụ', 'Loại cấu hình', 'Test bằng lệnh'),
        ('1', 'Tổng quỹ lương theo phòng ban', 'Calculated Field + Relationship', '"Tổng quỹ lương từng phòng ban?"'),
        ('2', 'Nhân viên > 10 năm chưa thăng chức', 'SQL Pair + Instruction', '"Nhân viên nào 10+ năm chưa thăng chức?"'),
        ('3', 'So sánh nghỉ việc: làm thêm giờ vs không', 'SQL Pair (GROUP BY)', '"So sánh tỷ lệ nghỉ việc OT vs không OT?"'),
        ('4', 'Nhân viên lương bất thường', 'Instruction (Global)', '"Nhân viên nào có lương bất thường?"'),
        ('5', 'Báo cáo tổng hợp', 'SQL Pair tổng hợp', '"Tổng NV, số rủi ro, tỷ lệ churn?"')
    ]
    
    for idx, row in enumerate(khải_uc):
        cells = khải_uc_table.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            for cell_idx, value in enumerate(row):
                cells[cell_idx].text = value
    
    # Hân
    add_heading_with_color(doc, '4.2 HÂN — 5 Nghiệp Vụ Semantic Layer', level=2, color=(204, 102, 0))
    
    hân_uc_table = doc.add_table(rows=6, cols=4)
    hân_uc_table.style = 'Light Grid Accent 1'
    add_table_border(hân_uc_table)
    
    hân_uc = [
        ('STT', 'Nghiệp vụ', 'Loại cấu hình', 'Test bằng lệnh'),
        ('1', 'Phân nhóm nhân viên theo tuổi', 'Calculated Field (Age_Group)', '"Phân nhóm NV theo tuổi?"'),
        ('2', 'Lương so với trung bình phòng', 'Relationship + Calculated', '"Ai có lương thấp hơn TB phòng?"'),
        ('3', 'Nhân viên undervalued', 'Instruction (Global)', '"NV nào có kinh nghiệm cao level thấp?"'),
        ('4', 'Work-life balance vs attrition', 'SQL Pair (multi-table)', '"Mối quan hệ work-life & rủi ro?"'),
        ('5', 'High Performer at Risk', 'SQL Pair (performance >=3)', '"High performer có rủi ro gì?"')
    ]
    
    for idx, row in enumerate(hân_uc):
        cells = hân_uc_table.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            for cell_idx, value in enumerate(row):
                cells[cell_idx].text = value
    
    # Ninh
    add_heading_with_color(doc, '4.3 NINH — 5 Nghiệp Vụ Agentic Layer', level=2, color=(204, 102, 0))
    
    ninh_uc_table = doc.add_table(rows=6, cols=4)
    ninh_uc_table.style = 'Light Grid Accent 1'
    add_table_border(ninh_uc_table)
    
    ninh_uc = [
        ('STT', 'Nghiệp vụ', 'Loại cấu hình', 'Test bằng lệnh'),
        ('1', 'Top 5 nhân viên rủi ro cao nhất', 'SQL Pair (CTE+ROW_NUMBER)', '"Top 5 NV rủi ro cao nhất?"'),
        ('2', 'Nhân viên có burnout nguy hiểm', 'SQL Pair + Instruction', '"Nhân viên nào có burnout nguy hiểm?"'),
        ('3', 'Lọc mặc định risk_level High/Critical', 'Instruction (Global)', '"Liệt kê NV rủi ro"'),
        ('4', 'So sánh rủi ro by job_role', 'SQL Pair (multi-group)', '"So sánh tỷ lệ rủi ro các job role?"'),
        ('5', 'Mặc định monthly_income', 'Instruction (Global)', '"Tính lương trung bình?"')
    ]
    
    for idx, row in enumerate(ninh_uc):
        cells = ninh_uc_table.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            for cell_idx, value in enumerate(row):
                cells[cell_idx].text = value
    
    doc.add_page_break()
    
    # ===== PHẦN 5: GIT WORKFLOW =====
    add_heading_with_color(doc, '✅ PHẦN 5: GIT WORKFLOW — PUSH LÊN GITLAB', level=1)
    
    add_heading_with_color(doc, '5.1 Quy Trình Standard Git', level=2, color=(0, 102, 204))
    
    git_steps = [
        ('git branch -a', 'Kiểm tra branch (phải là hr_domain_research)'),
        ('git status', 'Xem tệp đã thay đổi'),
        ('git add khải/*', 'Add folder cá nhân (hoặc git add .)'),
        ('git commit -m "Khải: Deep dive Q1-Q8 + 5 use cases"', 'Commit với message rõ'),
        ('git push origin hr_domain_research', 'Push lên GitLab')
    ]
    
    for idx, (cmd, desc) in enumerate(git_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(cmd)
        run.font.name = 'Courier New'
        p.add_run(f' → {desc}')
    
    add_heading_with_color(doc, '5.2 Cấu Trúc Folder trên GitLab', level=2, color=(0, 102, 204))
    
    struct = '''hr-ai-project/
├── khải/
│   ├── Q1_Q8_deepdive.docx
│   ├── 5_use_cases_screenshots/
│   └── README.md
├── hân/
│   ├── Q1_Q9_semantic.docx
│   ├── 5_use_cases_screenshots/
│   └── README.md
├── ninh/
│   ├── Q1_Q9_agentic.docx
│   ├── 5_use_cases_screenshots/
│   └── README.md
├── gia/
│   ├── ML_Analysis.docx
│   └── Slides.pptx
├── uyên/
│   ├── Architecture.docx
│   └── 3_Diagrams/
└── HUONG_DAN_TEAM_DEV.docx'''
    
    p = doc.add_paragraph(struct)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    
    add_heading_with_color(doc, '5.3 Commit Message Format', level=2, color=(0, 102, 204))
    
    msg_format = '[Tên_người]: [Công_việc] - [Deliverables]'
    
    p = doc.add_paragraph(f'Format: {msg_format}')
    for run in p.runs:
        run.font.name = 'Courier New'
    
    doc.add_paragraph()
    
    examples = [
        'Khải: Deep dive Q1-Q8 + 5 use cases + architecture diagram',
        'Hân: Semantic layer Q1-Q9 + 5 calculated fields + indexing diagram',
        'Ninh: Agentic layer Q1-Q9 + 5 SQL pairs + 5 instructions',
        'Gia: ML analysis + feature importance + business insights + slides',
        'Uyên: Project architecture + 3 diagrams + 12 Q answers'
    ]
    
    for ex in examples:
        p = doc.add_paragraph(ex, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== PHẦN 6: TIMELINE =====
    add_heading_with_color(doc, '✅ PHẦN 6: TIMELINE & DEADLINE', level=1)
    
    timeline_table = doc.add_table(rows=8, cols=4)
    timeline_table.style = 'Light Grid Accent 1'
    add_table_border(timeline_table)
    
    timeline_data = [
        ('Ngày', 'Hoạt động', 'Ai', 'Checklist'),
        ('11/02', 'Clone GitLab + Restore SQL + Chạy Docker', 'Tất cả', '☐ Repo UP + SQL connected'),
        ('12/02', 'Đọc source code + Ghi chú', 'Tất cả', '☐ Notes hoàn chỉnh'),
        ('13/02', 'Cấu hình 5 use cases + Test', 'Khải,Hân,Ninh', '☐ 5 use cases active'),
        ('13/02', 'Giải thích notebook + Biểu đồ', 'Gia', '☐ Draft report'),
        ('13/02', 'Vẽ sơ đồ + Viết tổng quan', 'Uyên', '☐ Draft diagrams'),
        ('14-15/02', 'Hoàn thiện + Review + Screenshot', 'Tất cả', '☐ Docx + Slides final'),
        ('16/02', 'Push GitLab + Submit', 'Tất cả', '☐ Push successful')
    ]
    
    for idx, row in enumerate(timeline_data):
        cells = timeline_table.rows[idx].cells
        if idx == 0:
            for cell in cells:
                set_cell_background(cell, '003366')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            for cell_idx, value in enumerate(row):
                cells[cell_idx].text = value
    
    doc.add_page_break()
    
    # ===== PHẦN 7: REFERENCE =====
    add_heading_with_color(doc, '✅ PHẦN 7: LIÊN HỆ & REFERENCE', level=1)
    
    add_heading_with_color(doc, '7.1 URLs & Links', level=2, color=(0, 102, 204))
    
    links = [
        ('GitLab Repository', 'https://gitlab.com/boygia757-netizen/hr-ai-project'),
        ('Public Demo', 'https://certification-lows-spy-tension.trycloudflare.com'),
        ('Local UI', 'http://localhost:3000'),
        ('AI Service Health', 'http://localhost:5555/health'),
        ('Qdrant Dashboard', 'http://localhost:6333/dashboard'),
        ('SQL Server (local)', 'localhost,1433 (SA / YourPassword)')
    ]
    
    for label, url in links:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        if 'http' in url:
            run = p.add_run(url)
            run.font.color.rgb = RGBColor(0, 0, 255)
        else:
            p.add_run(url)
    
    doc.add_paragraph()
    
    add_heading_with_color(doc, '7.2 Key Database Objects', level=2, color=(0, 102, 204))
    
    objects = [
        ('Table', 'ms_employee', 'Thông tin nhân viên'),
        ('Table', 'attrition_forecast', 'Dự báo rủi ro'),
        ('Table', 'hr_training_data', 'Dữ liệu training ML'),
        ('View', 'v_employee_actionable_insights', 'Insights có thể hành động'),
        ('Function', 'CalculateAttritionRisk()', 'Tính rủi ro trong real-time'),
    ]
    
    for obj_type, obj_name, desc in objects:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{obj_type}: ').bold = True
        p.add_run(f'{obj_name} → {desc}')
    
    doc.add_paragraph()
    
    add_heading_with_color(doc, '7.3 File quan trọng trong Repo', level=2, color=(0, 102, 204))
    
    important_files = [
        'HUONG_DAN_TEAM_DEV.docx — Hướng dẫn chính',
        'HUONG_DAN_CLONE_GITLAB_VA_SETUP.md — Hướng dẫn chi tiết',
        'HR_Analytics.bak — Backup SQL Server',
        'WrenAI/docker/docker-compose.yaml — Config 6 containers',
        'WrenAI/docker/.env — Environment variables',
        'notebooks/HR_Analytics_Project_Final.ipynb — ML Notebook',
        'legacy/init-db.sql — Init database schema'
    ]
    
    for file in important_files:
        doc.add_paragraph(file, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== FOOTER =====
    footer_text = f'Tài liệu được tạo: {datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")}'
    p = doc.add_paragraph(footer_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    p = doc.add_paragraph('Hướng dẫn đầy đủ cho team phát triển HR Analytics AI Project')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    return doc

if __name__ == '__main__':
    doc = create_detailed_guide_docx()
    output_path = 'HUONG_DAN_CHI_TIET_GITLAB_SQL_USE_CASES.docx'
    doc.save(output_path)
    print(f'✓ Tài liệu chi tiết đã tạo: {output_path}')
    print(f'✓ Nội dung: GitLab Access, SQL Restore, Docker, 5 Use Cases, Git Workflow, Timeline')
    print(f'✓ Dung lượng: ~50 KB')
